"""
File generation utilities for creating book files in different formats.
"""

import os
import tempfile
from typing import Optional
from datetime import datetime
import structlog
from flask import current_app

logger = structlog.get_logger()


def generate_pdf(book) -> Optional[str]:
    """
    Generate PDF file from book content.
    
    Args:
        book: BookGeneration instance
        
    Returns:
        Path to generated PDF file or None if failed
    """
    try:
        from reportlab.lib.pagesizes import letter, A4, A5, landscape
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        from reportlab.lib.units import inch, mm
        from reportlab.pdfgen import canvas
        
        # Create file path
        storage_dir = current_app.config.get('STORAGE_PATH', '/storage')
        books_dir = os.path.join(storage_dir, 'books')
        os.makedirs(books_dir, exist_ok=True)
        
        filename = f"book_{book.id}_{book.uuid}.pdf"
        file_path = os.path.join(books_dir, filename)
        
        # Determine page size
        page_sizes = {
            'pocket': (110*mm, 180*mm),  # Pocket size (como Kindle)
            'A5': A5,
            'B5': (176*mm, 250*mm),  # B5 custom size
            'letter': letter
        }
        page_size = page_sizes.get(book.format_size, letter)
        
        # Create PDF document
        doc = SimpleDocTemplate(file_path, pagesize=page_size)
        
        # Get styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'BookTitle',
            parent=styles['Title'],
            fontSize=24,
            textColor=colors.darkblue,
            spaceAfter=30,
            alignment=1  # Center
        )
        
        heading_style = ParagraphStyle(
            'BookHeading',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.darkblue,
            spaceAfter=12,
            spaceBefore=24
        )
        
        # Determine line spacing
        line_spacing_multipliers = {
            'single': 1.0,
            'medium': 1.5,
            'double': 2.0
        }
        line_spacing = line_spacing_multipliers.get(
            getattr(book, 'line_spacing', 'medium'), 1.5
        )
        
        normal_style = ParagraphStyle(
            'BookNormal',
            parent=styles['Normal'],
            fontSize=12,
            textColor=colors.black,
            spaceAfter=12 * line_spacing,
            leading=14 * line_spacing,  # Line height
            alignment=0  # Left
        )
        
        # Build content
        story = []
        
        # Title page
        story.append(Spacer(1, 2*inch))
        story.append(Paragraph(book.title, title_style))
        story.append(Spacer(1, 0.5*inch))
        story.append(Paragraph(f"Género: {book.genre or 'General'}", normal_style))
        story.append(Paragraph(f"Generado: {datetime.now().strftime('%d/%m/%Y')}", normal_style))
        story.append(PageBreak())
        
        # Process content
        if book.content:
            lines = book.content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    story.append(Spacer(1, 6))
                elif line.startswith('# '):
                    # Main title
                    story.append(Paragraph(line[2:], title_style))
                elif line.startswith('## '):
                    # Chapter/section heading
                    story.append(Paragraph(line[3:], heading_style))
                elif line.startswith('### '):
                    # Sub-heading
                    story.append(Paragraph(line[4:], styles['Heading2']))
                else:
                    # Normal paragraph
                    story.append(Paragraph(line, normal_style))
        
        # Build PDF
        doc.build(story)
        
        logger.info("pdf_generated", book_id=book.id, file_path=file_path)
        return file_path
        
    except ImportError as e:
        logger.error("reportlab_not_available", error=str(e))
        return None
    except Exception as e:
        logger.error("pdf_generation_failed", book_id=book.id, error=str(e))
        return None


def generate_epub(book) -> Optional[str]:
    """
    Generate EPUB file from book content.
    
    Args:
        book: BookGeneration instance
        
    Returns:
        Path to generated EPUB file or None if failed
    """
    try:
        from ebooklib import epub
        
        # Create file path
        storage_dir = current_app.config.get('STORAGE_PATH', '/storage')
        books_dir = os.path.join(storage_dir, 'books')
        os.makedirs(books_dir, exist_ok=True)
        
        filename = f"book_{book.id}_{book.uuid}.epub"
        file_path = os.path.join(books_dir, filename)
        
        # Create EPUB book
        epub_book = epub.EpubBook()
        
        # Set metadata
        epub_book.set_identifier(str(book.uuid))
        epub_book.set_title(book.title)
        epub_book.set_language(book.language or 'es')
        epub_book.add_author('Buko AI')
        
        # Create chapters from content
        chapters = []
        if book.content:
            content_parts = book.content.split('## ')
            
            for i, part in enumerate(content_parts):
                if not part.strip():
                    continue
                
                lines = part.strip().split('\n')
                chapter_title = lines[0] if lines else f'Capítulo {i}'
                chapter_content = '\n'.join(lines[1:]) if len(lines) > 1 else ''
                
                # Create chapter
                chapter = epub.EpubHtml(
                    title=chapter_title,
                    file_name=f'chapter_{i}.xhtml',
                    lang=book.language or 'es'
                )
                
                # Format content as HTML
                html_content = f"""
                <html>
                <head>
                    <title>{chapter_title}</title>
                </head>
                <body>
                    <h1>{chapter_title}</h1>
                    {''.join(f'<p>{line}</p>' for line in chapter_content.split('\n') if line.strip())}
                </body>
                </html>
                """
                
                chapter.content = html_content
                epub_book.add_item(chapter)
                chapters.append(chapter)
        
        # Add default CSS
        default_css = epub.EpubItem(
            uid="default",
            file_name="style/default.css",
            media_type="text/css",
            content="""
            body {
                font-family: Arial, sans-serif;
                margin: 1em;
                line-height: 1.5;
            }
            h1 {
                color: #333;
                border-bottom: 2px solid #333;
                padding-bottom: 0.5em;
            }
            p {
                margin-bottom: 1em;
                text-align: justify;
            }
            """
        )
        epub_book.add_item(default_css)
        
        # Create table of contents
        epub_book.toc = chapters
        
        # Add navigation files
        epub_book.add_item(epub.EpubNcx())
        epub_book.add_item(epub.EpubNav())
        
        # Create spine
        epub_book.spine = ['nav'] + chapters
        
        # Write EPUB file
        epub.write_epub(file_path, epub_book, {})
        
        logger.info("epub_generated", book_id=book.id, file_path=file_path)
        return file_path
        
    except ImportError as e:
        logger.error("ebooklib_not_available", error=str(e))
        return None
    except Exception as e:
        logger.error("epub_generation_failed", book_id=book.id, error=str(e))
        return None


def generate_docx(book) -> Optional[str]:
    """
    Generate DOCX file from book content.
    
    Args:
        book: BookGeneration instance
        
    Returns:
        Path to generated DOCX file or None if failed
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Create file path
        storage_dir = current_app.config.get('STORAGE_PATH', '/storage')
        books_dir = os.path.join(storage_dir, 'books')
        os.makedirs(books_dir, exist_ok=True)
        
        filename = f"book_{book.id}_{book.uuid}.docx"
        file_path = os.path.join(books_dir, filename)
        
        # Create document
        doc = Document()
        
        # Set document margins based on page size
        page_margins = {
            'pocket': (0.5, 0.5, 0.5, 0.5),  # Minimal margins for pocket
            'A5': (0.75, 0.75, 0.75, 0.75),  # Smaller margins for A5
            'B5': (0.85, 0.85, 0.85, 0.85),
            'letter': (1, 1, 1, 1)
        }
        margins = page_margins.get(book.format_size, (1, 1, 1, 1))
        
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(margins[0])
            section.bottom_margin = Inches(margins[1])
            section.left_margin = Inches(margins[2])
            section.right_margin = Inches(margins[3])
            
            # Set page size
            if book.format_size == 'pocket':
                section.page_width = Inches(4.33)  # 110mm
                section.page_height = Inches(7.09)  # 180mm
            elif book.format_size == 'A5':
                section.page_width = Inches(5.83)  # 148mm
                section.page_height = Inches(8.27)  # 210mm
            elif book.format_size == 'B5':
                section.page_width = Inches(6.93)  # 176mm
                section.page_height = Inches(9.84)  # 250mm
            # Letter is default size
        
        # Add title page
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(book.title)
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        doc.add_paragraph(f"Género: {book.genre or 'General'}")
        doc.add_paragraph(f"Generado: {datetime.now().strftime('%d/%m/%Y')}")
        doc.add_page_break()
        
        # Process content
        if book.content:
            lines = book.content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    doc.add_paragraph()
                elif line.startswith('# '):
                    # Main title
                    p = doc.add_paragraph()
                    run = p.add_run(line[2:])
                    run.font.size = Pt(20)
                    run.font.bold = True
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif line.startswith('## '):
                    # Chapter/section heading
                    p = doc.add_paragraph()
                    run = p.add_run(line[3:])
                    run.font.size = Pt(16)
                    run.font.bold = True
                elif line.startswith('### '):
                    # Sub-heading
                    p = doc.add_paragraph()
                    run = p.add_run(line[4:])
                    run.font.size = Pt(14)
                    run.font.bold = True
                else:
                    # Normal paragraph
                    p = doc.add_paragraph(line)
                    # Set line spacing
                    if hasattr(book, 'line_spacing'):
                        from docx.enum.text import WD_LINE_SPACING
                        if book.line_spacing == 'single':
                            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                        elif book.line_spacing == 'double':
                            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                        else:  # medium (1.5)
                            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        # Save document
        doc.save(file_path)
        
        logger.info("docx_generated", book_id=book.id, file_path=file_path)
        return file_path
        
    except ImportError as e:
        logger.error("python_docx_not_available", error=str(e))
        return None
    except Exception as e:
        logger.error("docx_generation_failed", book_id=book.id, error=str(e))
        return None


def generate_txt(book) -> Optional[str]:
    """
    Generate TXT file from book content.
    
    Args:
        book: BookGeneration instance
        
    Returns:
        Path to generated TXT file or None if failed
    """
    try:
        # Create file path
        storage_dir = current_app.config.get('STORAGE_PATH', '/storage')
        books_dir = os.path.join(storage_dir, 'books')
        os.makedirs(books_dir, exist_ok=True)
        
        filename = f"book_{book.id}_{book.uuid}.txt"
        file_path = os.path.join(books_dir, filename)
        
        # Create content
        content_lines = []
        
        # Add header
        content_lines.extend([
            "=" * 50,
            book.title.upper(),
            "=" * 50,
            "",
            f"Género: {book.genre or 'General'}",
            f"Generado: {datetime.now().strftime('%d/%m/%Y')}",
            "",
            "=" * 50,
            ""
        ])
        
        # Add book content
        if book.content:
            # Clean up markdown formatting for plain text
            lines = book.content.split('\n')
            for line in lines:
                line = line.strip()
                if line.startswith('# '):
                    content_lines.extend(["", "=" * 50, line[2:].upper(), "=" * 50, ""])
                elif line.startswith('## '):
                    content_lines.extend(["", "-" * 30, line[3:], "-" * 30, ""])
                elif line.startswith('### '):
                    content_lines.extend(["", line[4:], ""])
                else:
                    content_lines.append(line)
        
        # Write file
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(content_lines))
        
        logger.info("txt_generated", book_id=book.id, file_path=file_path)
        return file_path
        
    except Exception as e:
        logger.error("txt_generation_failed", book_id=book.id, error=str(e))
        return None


def cleanup_old_files(days: int = 30) -> int:
    """
    Clean up old book files.
    
    Args:
        days: Number of days to keep files
        
    Returns:
        Number of files cleaned up
    """
    try:
        import time
        
        storage_dir = current_app.config.get('STORAGE_PATH', '/storage')
        books_dir = os.path.join(storage_dir, 'books')
        
        if not os.path.exists(books_dir):
            return 0
        
        cutoff_time = time.time() - (days * 24 * 60 * 60)
        cleaned_count = 0
        
        for filename in os.listdir(books_dir):
            file_path = os.path.join(books_dir, filename)
            
            if os.path.isfile(file_path):
                file_mtime = os.path.getmtime(file_path)
                
                if file_mtime < cutoff_time:
                    try:
                        os.remove(file_path)
                        cleaned_count += 1
                        logger.info("old_file_cleaned", file_path=file_path)
                    except Exception as e:
                        logger.error("file_cleanup_failed", 
                                   file_path=file_path,
                                   error=str(e))
        
        logger.info("file_cleanup_completed", cleaned_count=cleaned_count)
        return cleaned_count
        
    except Exception as e:
        logger.error("file_cleanup_error", error=str(e))
        return 0