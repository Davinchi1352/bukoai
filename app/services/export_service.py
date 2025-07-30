"""
Export Service - Professional multi-format book export with platform-specific formatting.

Supports standard and platform-specific exports for:
- Amazon KDP
- Google Play Books
- Apple Books
- Kobo Writing Life
- Smashwords
- Gumroad
- Payhip
"""

import os
import io
import uuid
import mimetypes
from datetime import datetime
from typing import Optional, Dict, Any, Tuple, List
from enum import Enum
import structlog
from flask import current_app, send_file
from PIL import Image, ImageDraw, ImageFont
import textwrap

# Import required libraries
try:
    from reportlab.lib.pagesizes import letter, A4, A5, landscape
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle, Image as RLImage
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.units import inch, mm, cm
    from reportlab.pdfgen import canvas
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from ebooklib import epub
    EBOOKLIB_AVAILABLE = True
except ImportError:
    EBOOKLIB_AVAILABLE = False

# Import professional formatting services
try:
    from .professional_formatting_service import ProfessionalFormattingService, ProfessionalFormattingOptions
    from .book_formatting_service import FormattingPlatform, FormattingOptions
    from .markdown_to_html_service import BookStructure
    PROFESSIONAL_FORMATTING_AVAILABLE = True
except ImportError:
    PROFESSIONAL_FORMATTING_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.section import WD_SECTION
    from docx.oxml.shared import OxmlElement, qn
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

logger = structlog.get_logger()


class ExportFormat(Enum):
    """Supported export formats."""
    PDF = "pdf"
    EPUB = "epub"
    DOCX = "docx"
    TXT = "txt"


class ExportPlatform(Enum):
    """Supported publishing platforms."""
    STANDARD = "standard"
    AMAZON_KDP = "amazon_kdp"
    GOOGLE_PLAY = "google_play"
    APPLE_BOOKS = "apple_books"
    KOBO = "kobo"
    SMASHWORDS = "smashwords"
    GUMROAD = "gumroad"
    PAYHIP = "payhip"


class PlatformConfig:
    """Platform-specific configuration for professional book formatting following 2025 industry standards."""
    
    CONFIGS = {
        ExportPlatform.STANDARD: {
            "page_size": (15.24 * cm, 22.86 * cm),  # 6"x9" professional standard
            "margins": {"top": 2.54 * cm, "bottom": 2.54 * cm, "left": 2.54 * cm, "right": 2.54 * cm},  # 1" all around
            "fonts": {
                "body": "Garamond",  # Professional serif for body text
                "body_fallback": "Times New Roman",
                "heading": "Arial",  # Clean sans-serif for headings
                "title": "Arial",
                "subtitle": "Arial",
                "chapter": "Arial"
            },
            "font_sizes": {
                "body": 11,  # Professional body text size
                "heading": 14,  # Section headings
                "subheading": 12,  # Subsection headings  
                "title": 20,  # Main titles
                "subtitle": 16,  # Subtitles
                "chapter": 18,  # Chapter titles
                "toc": 12,  # Table of contents
                "metadata": 10  # Small metadata text
            },
            "line_spacing": 1.4,  # Professional line spacing
            "paragraph_spacing": {
                "before": 6,  # Points before paragraph
                "after": 6    # Points after paragraph
            },
            "heading_spacing": {
                "chapter_before": 36,
                "chapter_after": 24,
                "section_before": 18,
                "section_after": 12,
                "subsection_before": 12,
                "subsection_after": 6
            },
            "include_page_numbers": True,
            "include_headers": True, 
            "cover_size": (2560, 1600),  # 1.6:1 ratio
            "toc_required": True,
            "toc_interactive": True,
            "professional_formatting": True,
            "use_drop_caps": False,  # First letter enhancement
            "chapter_breaks": True,  # New page for each chapter
            "section_breaks": False,  # Just spacing for sections
            # Commercial ebook elements
            "include_copyright_page": True,
            "include_about_author": True,
            "include_disclaimer": True,
            "include_contact_info": True,
            "include_license": True,
            "page_numbering": True,
            "headers_footers": True,
            "professional_cover": True,
            "brand_elements": {
                "publisher": "Buko AI Editorial",
                "brand_color": "#2C5282",  # Professional blue
                "website": "https://buko.ai",
                "contact_email": "books@buko.ai"
            }
        },
        ExportPlatform.AMAZON_KDP: {
            "page_size": (15.24 * cm, 22.86 * cm),  # 6"x9"
            "margins": {"top": 2.5 * cm, "bottom": 2.5 * cm, "left": 2 * cm, "right": 2 * cm},
            "fonts": {
                "body": "Times New Roman",
                "heading": "Georgia",
                "title": "Georgia"
            },
            "font_sizes": {
                "body": 12,
                "heading": 16,
                "title": 24,
                "chapter": 18
            },
            "line_spacing": 1.5,
            "include_page_numbers": False,  # Kindle auto-generates
            "include_headers": False,  # Kindle ignores
            "cover_size": (2560, 1600),  # 1.6:1 ratio
            "toc_required": True,
            "toc_interactive": True
        },
        ExportPlatform.GOOGLE_PLAY: {
            "page_size": (15.24 * cm, 22.86 * cm),
            "margins": {"top": 2.5 * cm, "bottom": 2.5 * cm, "left": 2 * cm, "right": 2 * cm},
            "fonts": {
                "body": "Times New Roman",
                "heading": "Georgia",
                "title": "Georgia"
            },
            "font_sizes": {
                "body": 11,
                "heading": 16,
                "title": 24,
                "chapter": 18
            },
            "line_spacing": 1.15,
            "include_page_numbers": False,
            "include_headers": False,
            "cover_size": (1400, 2100),  # Min 1400px width
            "toc_required": True,
            "toc_interactive": True,
            "include_metadata": True
        },
        ExportPlatform.APPLE_BOOKS: {
            "page_size": (15.24 * cm, 22.86 * cm),
            "margins": {"top": 2.5 * cm, "bottom": 2.5 * cm, "left": 2 * cm, "right": 2 * cm},
            "fonts": {
                "body": "Georgia",
                "heading": "Georgia",
                "title": "Georgia"
            },
            "font_sizes": {
                "body": 12,
                "heading": 16,
                "title": 24,
                "chapter": 18
            },
            "line_spacing": 1.2,
            "include_page_numbers": False,  # Auto in app
            "include_headers": False,
            "cover_size": (1400, 2100),  # Min 1400px width
            "toc_required": True,
            "toc_interactive": True
        },
        ExportPlatform.KOBO: {
            "page_size": (15.24 * cm, 22.86 * cm),
            "margins": {"top": 2.5 * cm, "bottom": 2.5 * cm, "left": 2 * cm, "right": 2 * cm},
            "fonts": {
                "body": "Times New Roman",
                "heading": "Georgia",
                "title": "Georgia"
            },
            "font_sizes": {
                "body": 11,
                "heading": 16,
                "title": 24,
                "chapter": 18
            },
            "line_spacing": 1.15,
            "include_page_numbers": False,
            "include_headers": False,
            "cover_size": (1600, 2400),  # 2:3 ratio
            "toc_required": True,
            "toc_interactive": True
        },
        ExportPlatform.SMASHWORDS: {
            "page_size": letter,  # US Letter
            "margins": {"top": 2.5 * cm, "bottom": 2.5 * cm, "left": 2.5 * cm, "right": 2.5 * cm},
            "fonts": {
                "body": "Times New Roman",
                "heading": "Times New Roman",
                "title": "Times New Roman"
            },
            "font_sizes": {
                "body": 12,
                "heading": 16,
                "title": 24,
                "chapter": 18
            },
            "line_spacing": 1.5,
            "include_page_numbers": False,
            "include_headers": False,
            "cover_size": (1600, 2400),
            "toc_required": True,
            "toc_interactive": True,
            "use_simple_styles": True  # Smashwords prefers simple formatting
        },
        ExportPlatform.GUMROAD: {
            "page_size": (15.24 * cm, 22.86 * cm),  # 6"x9"
            "margins": {"top": 2 * cm, "bottom": 2 * cm, "left": 2 * cm, "right": 2 * cm},
            "fonts": {
                "body": "Arial",
                "heading": "Georgia",
                "title": "Georgia"
            },
            "font_sizes": {
                "body": 11,
                "heading": 16,
                "title": 24,
                "chapter": 18
            },
            "line_spacing": 1.15,
            "include_page_numbers": True,
            "include_headers": True,
            "cover_size": (1600, 2400),  # 2:3 ratio
            "toc_required": True,
            "toc_interactive": True,
            "include_copyright": True
        },
        ExportPlatform.PAYHIP: {
            "page_size": A4,
            "margins": {"top": 2 * cm, "bottom": 2 * cm, "left": 2 * cm, "right": 2 * cm},
            "fonts": {
                "body": "Arial",
                "heading": "Georgia",
                "title": "Georgia"
            },
            "font_sizes": {
                "body": 11,
                "heading": 16,
                "title": 24,
                "chapter": 18
            },
            "line_spacing": 1.2,
            "include_page_numbers": True,
            "include_headers": True,
            "cover_size": (1600, 2400),  # 2:3 ratio
            "toc_required": True,
            "toc_interactive": True,
            "include_metadata": True
        }
    }


class BookExportService:
    """Service for exporting books in multiple formats with platform-specific formatting."""
    
    def __init__(self):
        self.storage_dir = current_app.config.get('STORAGE_PATH', 'storage')
        self.books_dir = os.path.join(self.storage_dir, 'books')
        self.covers_dir = os.path.join(self.storage_dir, 'covers')
        os.makedirs(self.books_dir, exist_ok=True)
        os.makedirs(self.covers_dir, exist_ok=True)
        
        # Initialize professional formatting service if available
        if PROFESSIONAL_FORMATTING_AVAILABLE:
            self.professional_service = ProfessionalFormattingService()
        else:
            self.professional_service = None
    
    def export_book(self, book, format: ExportFormat, platform: ExportPlatform = ExportPlatform.STANDARD) -> Optional[str]:
        """
        Export a book in the specified format for the specified platform.
        
        Args:
            book: BookGeneration instance
            format: Export format (PDF, EPUB, DOCX)
            platform: Target platform for formatting
            
        Returns:
            Path to generated file or None if failed
        """
        try:
            logger.info("book_export_started", 
                       book_id=book.id, 
                       format=format.value, 
                       platform=platform.value)
            
            # Get platform configuration
            config = PlatformConfig.CONFIGS.get(platform, PlatformConfig.CONFIGS[ExportPlatform.STANDARD])
            
            # Create professional formatting options if service is available
            if self.professional_service:
                professional_options = self._create_professional_options(config, platform)
                formatted_result = self._get_professional_formatted_content(book, professional_options)
            else:
                formatted_result = None
                professional_options = None
            
            # Generate cover if needed
            cover_path = self._generate_cover(book, config)
            
            # Export based on format
            if format == ExportFormat.PDF:
                return self._export_pdf(book, config, cover_path, formatted_result)
            elif format == ExportFormat.EPUB:
                return self._export_epub(book, config, cover_path, formatted_result)
            elif format == ExportFormat.DOCX:
                return self._export_docx(book, config, cover_path, formatted_result)
            elif format == ExportFormat.TXT:
                return self._export_txt(book)
            else:
                raise ValueError(f"Unsupported format: {format}")
                
        except Exception as e:
            logger.error("book_export_failed", 
                        book_id=book.id, 
                        format=format.value,
                        platform=platform.value,
                        error=str(e))
            return None
    
    def _generate_cover(self, book, config: Dict[str, Any]) -> Optional[str]:
        """Generate a professional book cover."""
        try:
            width, height = config['cover_size']
            
            # Create cover image
            cover = Image.new('RGB', (width, height), color='white')
            draw = ImageDraw.Draw(cover)
            
            # Create sophisticated multi-layer gradient based on genre
            for y in range(height):
                ratio = y / height
                
                # Genre-specific professional color schemes
                if book.genre and any(word in book.genre.lower() for word in ['business', 'negocio', 'empresa']):
                    # Professional navy to silver gradient
                    r = int(25 + (120 - 25) * ratio)
                    g = int(55 + (140 - 55) * ratio)
                    b = int(95 + (160 - 95) * ratio)
                elif book.genre and any(word in book.genre.lower() for word in ['health', 'medicina', 'salud']):
                    # Medical green gradient
                    r = int(20 + (70 - 20) * ratio)
                    g = int(80 + (150 - 80) * ratio)
                    b = int(40 + (110 - 40) * ratio)
                elif book.genre and any(word in book.genre.lower() for word in ['tech', 'programming', 'tecnologia']):
                    # Tech purple gradient
                    r = int(60 + (130 - 60) * ratio)
                    g = int(30 + (80 - 30) * ratio)
                    b = int(120 + (180 - 120) * ratio)
                elif book.genre and any(word in book.genre.lower() for word in ['education', 'educacion', 'self_help']):
                    # Educational blue gradient  
                    r = int(30 + (90 - 30) * ratio)
                    g = int(60 + (140 - 60) * ratio)
                    b = int(120 + (200 - 120) * ratio)
                else:
                    # Default professional gradient
                    r = int(40 + (100 - 40) * ratio)
                    g = int(70 + (130 - 70) * ratio)
                    b = int(110 + (180 - 110) * ratio)
                
                draw.line([(0, y), (width, y)], fill=(r, g, b))
            
            # Add subtle texture overlay for professional appearance
            for i in range(0, width, 30):
                for j in range(0, height, 30):
                    if (i + j) % 60 == 0:
                        draw.ellipse([i, j, i+3, j+3], fill=(255, 255, 255, 20))
            
            # Load professional fonts with extensive fallback system
            try:
                # Extended font paths for better typography
                font_paths = [
                    # Linux fonts
                    "/usr/share/fonts/truetype/liberation/LiberationSerif-Bold.ttf",
                    "/usr/share/fonts/truetype/dejavu/DejaVuSerif-Bold.ttf", 
                    "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf",
                    "/usr/share/fonts/TTF/arial.ttf",
                    # macOS fonts
                    "/System/Library/Fonts/Helvetica.ttc",
                    "/System/Library/Fonts/Times.ttc",
                    # Windows fonts
                    "C:\\Windows\\Fonts\\arial.ttf",
                    "C:\\Windows\\Fonts\\times.ttf",
                ]
                title_font = None
                subtitle_font = None
                author_font = None
                
                for font_path in font_paths:
                    if os.path.exists(font_path):
                        try:
                            title_font = ImageFont.truetype(font_path, int(height * 0.09))
                            subtitle_font = ImageFont.truetype(font_path, int(height * 0.05))
                            author_font = ImageFont.truetype(font_path, int(height * 0.035))
                            break
                        except:
                            continue
                
                if not title_font:
                    title_font = ImageFont.load_default()
                    subtitle_font = ImageFont.load_default()
                    author_font = ImageFont.load_default()
                    
            except Exception as e:
                logger.warning("professional_font_loading_failed", error=str(e))
                title_font = ImageFont.load_default()
                subtitle_font = ImageFont.load_default()
                author_font = ImageFont.load_default()
            
            # Professional title processing
            title_parts = book.title.split(':')
            main_title = title_parts[0].strip().upper()
            subtitle_text = title_parts[1].strip() if len(title_parts) > 1 else ""
            
            # Smart title wrapping
            title_words = main_title.split()
            title_lines = []
            current_line = []
            
            for word in title_words:
                test_line = ' '.join(current_line + [word])
                if len(test_line) * 25 <= width - 200:  # Leave margins
                    current_line.append(word)
                else:
                    if current_line:
                        title_lines.append(' '.join(current_line))
                        current_line = [word]
                    else:
                        title_lines.append(word)
            
            if current_line:
                title_lines.append(' '.join(current_line))
            
            # Add professional title with shadow effects
            start_y = int(height * 0.25)
            current_y = start_y
            
            for line in title_lines[:3]:  # Limit to 3 lines
                try:
                    bbox = draw.textbbox((0, 0), line, font=title_font)
                    text_width = bbox[2] - bbox[0]
                    text_height = bbox[3] - bbox[1]
                    x = (width - text_width) // 2
                    
                    # Multiple shadow layers for professional depth
                    draw.text((x + 5, current_y + 5), line, fill=(0, 0, 0, 200), font=title_font)
                    draw.text((x + 3, current_y + 3), line, fill=(0, 0, 0, 120), font=title_font)
                    draw.text((x + 1, current_y + 1), line, fill=(0, 0, 0, 60), font=title_font)
                    draw.text((x, current_y), line, fill='white', font=title_font)
                    
                    current_y += text_height + 15
                except:
                    # Fallback for bbox issues
                    x = width // 4
                    draw.text((x, current_y), line, fill='white', font=title_font)
                    current_y += 60
            
            # Add subtitle with elegant styling
            if subtitle_text:
                current_y += 30
                try:
                    bbox = draw.textbbox((0, 0), subtitle_text, font=subtitle_font)
                    text_width = bbox[2] - bbox[0]
                    x = (width - text_width) // 2
                    
                    draw.text((x + 2, current_y + 2), subtitle_text, fill=(0, 0, 0, 150), font=subtitle_font)
                    draw.text((x, current_y), subtitle_text, fill='white', font=subtitle_font)
                    current_y += bbox[3] - bbox[1] + 40
                except:
                    x = width // 4
                    draw.text((x, current_y), subtitle_text, fill='white', font=subtitle_font)
                    current_y += 50
            
            # Add professional genre badge
            if book.genre:
                genre_text = book.genre.upper()
                try:
                    genre_font = ImageFont.truetype(font_paths[0] if os.path.exists(font_paths[0]) else font_paths[-1], int(height * 0.03))
                except:
                    genre_font = author_font
                
                try:
                    bbox = draw.textbbox((0, 0), genre_text, font=genre_font)
                    text_width = bbox[2] - bbox[0]
                    text_height = bbox[3] - bbox[1]
                    
                    # Create rounded rectangle badge
                    badge_width = text_width + 60
                    badge_height = text_height + 20
                    badge_x = (width - badge_width) // 2
                    badge_y = current_y + 20
                    
                    # Badge background
                    badge_rect = [badge_x, badge_y, badge_x + badge_width, badge_y + badge_height]
                    draw.rounded_rectangle(badge_rect, radius=badge_height//2, fill=(255, 255, 255, 220))
                    
                    # Badge text
                    text_x = badge_x + 30
                    text_y = badge_y + 10
                    draw.text((text_x, text_y), genre_text, fill=(50, 50, 50), font=genre_font)
                except:
                    # Fallback
                    x = (width - len(genre_text) * 10) // 2
                    draw.text((x, current_y + 20), f"[ {genre_text} ]", fill='white', font=author_font)
            
            # Add Buko AI branding with professional styling
            publisher = config['brand_elements']['publisher'].upper()
            tagline = "INTELIGENCIA ARTIFICIAL EDUCATIVA"
            
            try:
                # Publisher name
                pub_bbox = draw.textbbox((0, 0), publisher, font=author_font)
                pub_width = pub_bbox[2] - pub_bbox[0]
                pub_x = (width - pub_width) // 2
                pub_y = height - 120
                
                draw.text((pub_x + 2, pub_y + 2), publisher, fill=(0, 0, 0, 180), font=author_font)
                draw.text((pub_x, pub_y), publisher, fill='white', font=author_font)
                
                # Tagline with smaller font
                try:
                    tagline_font = ImageFont.truetype(font_paths[0] if os.path.exists(font_paths[0]) else font_paths[-1], int(height * 0.025))
                except:
                    tagline_font = author_font
                
                tag_bbox = draw.textbbox((0, 0), tagline, font=tagline_font)
                tag_width = tag_bbox[2] - tag_bbox[0]
                tag_x = (width - tag_width) // 2
                tag_y = pub_y + pub_bbox[3] - pub_bbox[1] + 10
                
                draw.text((tag_x + 1, tag_y + 1), tagline, fill=(0, 0, 0, 120), font=tagline_font)
                draw.text((tag_x, tag_y), tagline, fill=(200, 200, 200), font=tagline_font)
                
            except:
                # Fallback positioning
                x = width // 4
                draw.text((x, height - 100), publisher, fill='white', font=author_font)
            
            # Add professional decorative elements
            # Top accent bar
            accent_height = 10
            draw.rectangle([(0, 0), (width, accent_height)], fill=(255, 255, 255, 120))
            
            # Corner flourishes
            corner_size = 80
            for i in range(corner_size):
                alpha = int(100 * (1 - i / corner_size))
                # Top corners
                draw.point((i, corner_size - i), fill=(255, 255, 255, alpha))
                draw.point((width - i - 1, corner_size - i), fill=(255, 255, 255, alpha))
            
            # Bottom accent lines
            for i in range(3):
                y = height - 160 - (i * 12)
                line_width = width // 4
                line_x = (width - line_width) // 2
                alpha = 140 - (i * 30)
                draw.line([(line_x, y), (line_x + line_width, y)], fill=(255, 255, 255, alpha), width=3)
            
            # Save cover with highest commercial quality
            filename = f"cover_{book.id}_{book.uuid}.jpg"
            cover_path = os.path.join(self.covers_dir, filename)
            cover.save(cover_path, 'JPEG', quality=98, optimize=True, progressive=True, dpi=(300, 300))
            
            logger.info("professional_cover_generated", 
                       book_id=book.id, 
                       cover_path=cover_path,
                       genre=book.genre,
                       publisher=config['brand_elements']['publisher'])
            return cover_path
            
        except Exception as e:
            logger.error("cover_generation_failed", book_id=book.id, error=str(e))
            return None
    
    def _create_professional_options(self, config: Dict[str, Any], platform: ExportPlatform) -> 'ProfessionalFormattingOptions':
        """Create professional formatting options from platform config."""
        if not PROFESSIONAL_FORMATTING_AVAILABLE:
            return None
            
        # Map platform to formatting platform
        platform_map = {
            ExportPlatform.AMAZON_KDP: FormattingPlatform.AMAZON_KDP,
            ExportPlatform.GOOGLE_PLAY: FormattingPlatform.GOOGLE_PLAY,
            ExportPlatform.APPLE_BOOKS: FormattingPlatform.APPLE_BOOKS,
            ExportPlatform.KOBO: FormattingPlatform.KOBO,
            ExportPlatform.STANDARD: FormattingPlatform.STANDARD,
        }
        
        formatting_platform = platform_map.get(platform, FormattingPlatform.STANDARD)
        
        return ProfessionalFormattingOptions(
            platform=formatting_platform,
            font_family=config['fonts']['body'],
            font_size_body=config['font_sizes']['body'],
            line_spacing=config['line_spacing'],
            use_professional_typography=True,
            use_chapter_breaks=config.get('chapter_breaks', True),
            use_headers_footers=config.get('include_headers', True),
            include_table_of_contents=config.get('toc_required', True),
            include_cover_page=True,
            include_title_page=config.get('professional_formatting', True),
            include_copyright_page=config.get('include_copyright_page', True),
            include_about_author=config.get('include_about_author', True),
            enable_toc_navigation=config.get('toc_interactive', True),
            enable_index_generation=True,
            enable_cross_references=True,
            enable_footnotes=True,
            enable_page_numbers=config.get('include_page_numbers', True),
            theme="professional",
            color_scheme="default",
            optimize_file_size=True,
            embed_fonts=True,
            include_metadata=config.get('include_metadata', True),
            enable_highlights=True,
            enable_annotations=True,
            enable_bookmarks=True,
            enable_search=True,
            include_isbn="",
            include_publisher_info=True,
            include_legal_notice=True,
            include_marketing_pages=False
        )
    
    def _get_professional_formatted_content(self, book, options: 'ProfessionalFormattingOptions') -> Optional[Dict[str, Any]]:
        """Get professionally formatted content using the new service."""
        if not self.professional_service or not options:
            return None
            
        try:
            return self.professional_service.format_for_commercial_distribution(book, options)
        except Exception as e:
            logger.warning("professional_formatting_failed", book_id=book.id, error=str(e))
            return None
    
    def _export_pdf(self, book, config: Dict[str, Any], cover_path: Optional[str], formatted_result: Optional[Dict[str, Any]] = None) -> Optional[str]:
        """Export book as PDF with platform-specific formatting."""
        if not REPORTLAB_AVAILABLE:
            logger.error("reportlab_not_available")
            return None
            
        try:
            filename = f"book_{book.id}_{book.uuid}.pdf"
            file_path = os.path.join(self.books_dir, filename)
            
            # Create PDF document with platform-specific page size
            doc = SimpleDocTemplate(
                file_path,
                pagesize=config['page_size'],
                topMargin=config['margins']['top'],
                bottomMargin=config['margins']['bottom'],
                leftMargin=config['margins']['left'],
                rightMargin=config['margins']['right']
            )
            
            # Create styles
            styles = self._create_pdf_styles(config)
            story = []
            
            # Add cover page if available
            if cover_path and os.path.exists(cover_path):
                try:
                    cover_width = config['page_size'][0] - config['margins']['left'] - config['margins']['right']
                    cover_height = cover_width * 1.5  # Maintain aspect ratio
                    story.append(RLImage(cover_path, width=cover_width, height=cover_height))
                    story.append(PageBreak())
                except Exception as e:
                    logger.warning("pdf_cover_failed", error=str(e))
            
            # Title page
            story.append(Spacer(1, 2*inch))
            story.append(Paragraph(book.title, styles['Title']))
            story.append(Spacer(1, 0.5*inch))
            
            # Metadata
            metadata = [
                f"Género: {book.genre or 'General'}",
                f"Idioma: {'Español' if book.language == 'es' else 'English' if book.language == 'en' else book.language}",
                f"Audiencia: {book.target_audience or 'General'}",
                f"Generado: {datetime.now().strftime('%d/%m/%Y')}"
            ]
            for meta in metadata:
                story.append(Paragraph(meta, styles['Metadata']))
                story.append(Spacer(1, 0.2*inch))
            
            # Copyright if required
            if config.get('include_copyright'):
                story.append(Spacer(1, 1*inch))
                story.append(Paragraph("© 2024 - Generado por Buko AI", styles['Copyright']))
            
            story.append(PageBreak())
            
            # Table of Contents if required
            if config.get('toc_required'):
                story.extend(self._create_pdf_toc(book, styles))
                story.append(PageBreak())
            
            # Process book content - use professional formatted content if available
            if formatted_result and formatted_result.get('structure'):
                story.extend(self._process_professional_content_for_pdf(formatted_result['structure'], styles))
            else:
                # Fallback to markdown content
                content = book.content_html if hasattr(book, 'content_html') and book.content_html else book.content
                story.extend(self._process_content_for_pdf(content, styles))
            
            # Build PDF
            doc.build(story)
            
            logger.info("pdf_exported", book_id=book.id, file_path=file_path)
            return file_path
            
        except Exception as e:
            logger.error("pdf_export_error", book_id=book.id, error=str(e))
            return None
    
    def _export_epub(self, book, config: Dict[str, Any], cover_path: Optional[str], formatted_result: Optional[Dict[str, Any]] = None) -> Optional[str]:
        """Export book as EPUB with platform-specific formatting."""
        if not EBOOKLIB_AVAILABLE:
            logger.error("ebooklib_not_available")
            return None
            
        try:
            filename = f"book_{book.id}_{book.uuid}.epub"
            file_path = os.path.join(self.books_dir, filename)
            
            # Create EPUB book
            epub_book = epub.EpubBook()
            
            # Set metadata
            epub_book.set_identifier(str(book.uuid))
            epub_book.set_title(book.title)
            epub_book.set_language(book.language or 'es')
            epub_book.add_author('Buko AI')
            
            if config.get('include_metadata'):
                epub_book.add_metadata('DC', 'subject', book.genre or 'General')
                epub_book.add_metadata('DC', 'description', f'Generated by Buko AI on {datetime.now().strftime("%Y-%m-%d")}')
            
            # Add cover if available
            if cover_path and os.path.exists(cover_path):
                try:
                    with open(cover_path, 'rb') as f:
                        epub_book.set_cover("cover.jpg", f.read())
                except Exception as e:
                    logger.warning("epub_cover_failed", error=str(e))
            
            # Create CSS for styling
            css_content = self._create_epub_css(config)
            css_item = epub.EpubItem(
                uid="style",
                file_name="style/main.css",
                media_type="text/css",
                content=css_content
            )
            epub_book.add_item(css_item)
            
            # Process content into chapters - use professional formatted content if available
            if formatted_result and formatted_result.get('structure'):
                chapters = self._process_professional_content_for_epub(formatted_result['structure'], config)
            else:
                chapters = self._process_content_for_epub(book, config)
            toc_entries = []
            
            for i, (chapter_title, chapter_content) in enumerate(chapters):
                chapter = epub.EpubHtml(
                    title=chapter_title,
                    file_name=f'chapter_{i+1}.xhtml',
                    lang=book.language or 'es'
                )
                chapter.content = chapter_content
                chapter.add_item(css_item)
                epub_book.add_item(chapter)
                toc_entries.append(chapter)
            
            # Create table of contents
            epub_book.toc = toc_entries
            
            # Add navigation files
            epub_book.add_item(epub.EpubNcx())
            epub_book.add_item(epub.EpubNav())
            
            # Create spine
            epub_book.spine = ['nav'] + toc_entries
            
            # Write EPUB file
            epub.write_epub(file_path, epub_book, {})
            
            logger.info("epub_exported", book_id=book.id, file_path=file_path)
            return file_path
            
        except Exception as e:
            logger.error("epub_export_error", book_id=book.id, error=str(e))
            return None
    
    def _export_docx(self, book, config: Dict[str, Any], cover_path: Optional[str], formatted_result: Optional[Dict[str, Any]] = None) -> Optional[str]:
        """Export book as DOCX with platform-specific formatting."""
        if not PYTHON_DOCX_AVAILABLE:
            logger.error("python_docx_not_available")
            return None
            
        try:
            filename = f"book_{book.id}_{book.uuid}.docx"
            file_path = os.path.join(self.books_dir, filename)
            
            # Create document
            doc = Document()
            
            # Set page size and margins
            section = doc.sections[0]
            
            # Set margins - config values are in reportlab points, convert to docx cm
            from reportlab.lib.units import cm as reportlab_cm
            section.top_margin = Cm(config['margins']['top'] / reportlab_cm)
            section.bottom_margin = Cm(config['margins']['bottom'] / reportlab_cm) 
            section.left_margin = Cm(config['margins']['left'] / reportlab_cm)
            section.right_margin = Cm(config['margins']['right'] / reportlab_cm)
            
            # Set page size based on platform
            if isinstance(config['page_size'], tuple):
                section.page_width = Cm(config['page_size'][0] / reportlab_cm)
                section.page_height = Cm(config['page_size'][1] / reportlab_cm)
            
            # Create professional title page
            self._create_professional_title_page(doc, book, config)
            
            # Process content with professional formatting - use professional formatted content if available
            if formatted_result and formatted_result.get('structure'):
                self._process_professional_content_for_docx(doc, formatted_result['structure'], config)
            else:
                # Fallback to markdown content
                content = book.content_html if hasattr(book, 'content_html') and book.content_html else book.content
                self._process_content_for_docx(doc, content, config)
            
            # Save document
            doc.save(file_path)
            
            logger.info("docx_exported", book_id=book.id, file_path=file_path)
            return file_path
            
        except Exception as e:
            logger.error("docx_export_error", book_id=book.id, error=str(e))
            return None
    
    def _export_txt(self, book) -> Optional[str]:
        """Export book as plain text."""
        try:
            filename = f"book_{book.id}_{book.uuid}.txt"
            file_path = os.path.join(self.books_dir, filename)
            
            content_lines = []
            
            # Header
            content_lines.extend([
                "=" * 80,
                book.title.upper().center(80),
                "=" * 80,
                "",
                f"Género: {book.genre or 'General'}",
                f"Generado: {datetime.now().strftime('%d/%m/%Y')}",
                f"Generado por: Buko AI",
                "",
                "=" * 80,
                "",
            ])
            
            # Process content
            if book.content:
                lines = book.content.split('\n')
                for line in lines:
                    line = line.strip()
                    if line.startswith('# '):
                        content_lines.extend(["", "=" * 80, line[2:].upper().center(80), "=" * 80, ""])
                    elif line.startswith('## '):
                        content_lines.extend(["", "-" * 60, line[3:].center(60), "-" * 60, ""])
                    elif line.startswith('### '):
                        content_lines.extend(["", line[4:], "-" * len(line[4:]), ""])
                    else:
                        content_lines.append(line)
            
            # Write file
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(content_lines))
            
            logger.info("txt_exported", book_id=book.id, file_path=file_path)
            return file_path
            
        except Exception as e:
            logger.error("txt_export_error", book_id=book.id, error=str(e))
            return None
    
    def _create_pdf_styles(self, config: Dict[str, Any]) -> Dict[str, ParagraphStyle]:
        """Create PDF paragraph styles based on platform configuration."""
        styles = getSampleStyleSheet()
        
        # Title style
        styles.add(ParagraphStyle(
            'BookTitle',
            parent=styles['Title'],
            fontName=config['fonts']['title'],
            fontSize=config['font_sizes']['title'],
            textColor=colors.HexColor('#1e3a8a'),
            spaceAfter=30,
            alignment=1  # Center
        ))
        
        # Chapter heading style
        styles.add(ParagraphStyle(
            'ChapterHeading',
            parent=styles['Heading1'],
            fontName=config['fonts']['heading'],
            fontSize=config['font_sizes']['chapter'],
            textColor=colors.HexColor('#1e3a8a'),
            spaceAfter=24,
            spaceBefore=36,
            keepWithNext=1
        ))
        
        # Section heading style
        styles.add(ParagraphStyle(
            'SectionHeading',
            parent=styles['Heading2'],
            fontName=config['fonts']['heading'],
            fontSize=config['font_sizes']['heading'],
            textColor=colors.HexColor('#3b82f6'),
            spaceAfter=18,
            spaceBefore=24
        ))
        
        # Body text style
        styles.add(ParagraphStyle(
            'BookBody',
            parent=styles['Normal'],
            fontName=config['fonts']['body'],
            fontSize=config['font_sizes']['body'],
            leading=config['font_sizes']['body'] * config['line_spacing'],
            textColor=colors.black,
            spaceAfter=12,
            alignment=0,  # Left aligned
            firstLineIndent=12 if not config.get('use_simple_styles') else 0
        ))
        
        # Metadata style
        styles.add(ParagraphStyle(
            'Metadata',
            parent=styles['Normal'],
            fontName=config['fonts']['body'],
            fontSize=config['font_sizes']['body'] - 2,
            textColor=colors.grey,
            alignment=1  # Center
        ))
        
        # Copyright style
        styles.add(ParagraphStyle(
            'Copyright',
            parent=styles['Normal'],
            fontName=config['fonts']['body'],
            fontSize=config['font_sizes']['body'] - 3,
            textColor=colors.grey,
            alignment=1  # Center
        ))
        
        # TOC style
        styles.add(ParagraphStyle(
            'TOCEntry',
            parent=styles['Normal'],
            fontName=config['fonts']['body'],
            fontSize=config['font_sizes']['body'],
            leftIndent=20,
            spaceAfter=8
        ))
        
        return styles
    
    def _create_pdf_toc(self, book, styles: Dict[str, ParagraphStyle]) -> List:
        """Create table of contents for PDF."""
        toc_elements = []
        
        # TOC Title
        toc_elements.append(Paragraph("Índice", styles['BookTitle']))
        toc_elements.append(Spacer(1, 0.5*inch))
        
        # Parse content to find chapters
        if book.content:
            lines = book.content.split('\n')
            chapter_num = 0
            
            for line in lines:
                line = line.strip()
                if line.startswith('## ') and not line.startswith('### '):
                    chapter_num += 1
                    chapter_title = line[3:].strip()
                    toc_entry = f"{chapter_num}. {chapter_title}"
                    toc_elements.append(Paragraph(toc_entry, styles['TOCEntry']))
        
        return toc_elements
    
    def _process_content_for_pdf(self, content: str, styles: Dict[str, ParagraphStyle]) -> List:
        """Process book content into PDF story elements."""
        story = []
        
        if not content:
            return story
        
        lines = content.split('\n')
        in_paragraph = False
        paragraph_text = []
        
        for line in lines:
            line = line.strip()
            
            if not line:
                # Empty line - end current paragraph if any
                if paragraph_text:
                    story.append(Paragraph(' '.join(paragraph_text), styles['BookBody']))
                    paragraph_text = []
                    in_paragraph = False
                story.append(Spacer(1, 6))
                
            elif line.startswith('# '):
                # Main title (skip, already in title page)
                continue
                
            elif line.startswith('## ') and not line.startswith('### '):
                # Chapter heading
                if paragraph_text:
                    story.append(Paragraph(' '.join(paragraph_text), styles['BookBody']))
                    paragraph_text = []
                    in_paragraph = False
                story.append(Paragraph(line[3:], styles['ChapterHeading']))
                
            elif line.startswith('### '):
                # Section heading
                if paragraph_text:
                    story.append(Paragraph(' '.join(paragraph_text), styles['BookBody']))
                    paragraph_text = []
                    in_paragraph = False
                story.append(Paragraph(line[4:], styles['SectionHeading']))
                
            else:
                # Regular paragraph text
                paragraph_text.append(line)
                in_paragraph = True
        
        # Don't forget the last paragraph
        if paragraph_text:
            story.append(Paragraph(' '.join(paragraph_text), styles['BookBody']))
        
        return story
    
    def _create_epub_css(self, config: Dict[str, Any]) -> str:
        """Create CSS for EPUB styling."""
        line_height = config['line_spacing'] * 100
        
        return f"""
        @page {{
            margin-top: {config['margins']['top'] / cm}cm;
            margin-bottom: {config['margins']['bottom'] / cm}cm;
            margin-left: {config['margins']['left'] / cm}cm;
            margin-right: {config['margins']['right'] / cm}cm;
        }}
        
        body {{
            font-family: {config['fonts']['body']}, serif;
            font-size: {config['font_sizes']['body']}pt;
            line-height: {line_height}%;
            color: #000;
            text-align: justify;
            margin: 0;
            padding: 0;
        }}
        
        h1 {{
            font-family: {config['fonts']['title']}, serif;
            font-size: {config['font_sizes']['title']}pt;
            color: #1e3a8a;
            text-align: center;
            margin: 2em 0 1em 0;
            page-break-before: always;
        }}
        
        h2 {{
            font-family: {config['fonts']['heading']}, serif;
            font-size: {config['font_sizes']['chapter']}pt;
            color: #1e3a8a;
            margin: 1.5em 0 1em 0;
            page-break-before: always;
        }}
        
        h3 {{
            font-family: {config['fonts']['heading']}, serif;
            font-size: {config['font_sizes']['heading']}pt;
            color: #3b82f6;
            margin: 1.2em 0 0.8em 0;
        }}
        
        p {{
            margin: 0 0 1em 0;
            text-indent: {"1.5em" if not config.get('use_simple_styles') else "0"};
        }}
        
        p:first-of-type {{
            text-indent: 0;
        }}
        
        .metadata {{
            text-align: center;
            color: #666;
            font-size: {config['font_sizes']['body'] - 2}pt;
            margin: 0.5em 0;
        }}
        
        .toc {{
            page-break-before: always;
        }}
        
        .toc-entry {{
            margin: 0.5em 0 0.5em 2em;
        }}
        """
    
    def _process_content_for_epub(self, book, config: Dict[str, Any]) -> List[Tuple[str, str]]:
        """Process book content into EPUB chapters."""
        chapters = []
        
        if not book.content:
            return chapters
        
        # Split content by chapters (## headings)
        content_parts = book.content.split('## ')
        
        # First part might be introduction or just the title
        if content_parts[0].strip():
            intro_content = self._format_epub_chapter("Introducción", content_parts[0].strip())
            if intro_content:
                chapters.append(("Introducción", intro_content))
        
        # Process each chapter
        for i, part in enumerate(content_parts[1:], 1):
            if not part.strip():
                continue
            
            lines = part.strip().split('\n')
            chapter_title = lines[0].strip() if lines else f'Capítulo {i}'
            chapter_content = '\n'.join(lines[1:]) if len(lines) > 1 else ''
            
            html_content = self._format_epub_chapter(chapter_title, chapter_content)
            chapters.append((chapter_title, html_content))
        
        return chapters
    
    def _format_epub_chapter(self, title: str, content: str) -> str:
        """Format a chapter for EPUB."""
        html_parts = [
            '<?xml version="1.0" encoding="UTF-8"?>',
            '<!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.1//EN" "http://www.w3.org/TR/xhtml11/DTD/xhtml11.dtd">',
            '<html xmlns="http://www.w3.org/1999/xhtml">',
            '<head>',
            f'<title>{title}</title>',
            '<link rel="stylesheet" type="text/css" href="style/main.css"/>',
            '</head>',
            '<body>',
            f'<h2>{title}</h2>'
        ]
        
        # Process content
        if content:
            lines = content.split('\n')
            in_paragraph = False
            paragraph_lines = []
            
            for line in lines:
                line = line.strip()
                
                if not line:
                    if paragraph_lines:
                        html_parts.append(f'<p>{" ".join(paragraph_lines)}</p>')
                        paragraph_lines = []
                        in_paragraph = False
                
                elif line.startswith('### '):
                    if paragraph_lines:
                        html_parts.append(f'<p>{" ".join(paragraph_lines)}</p>')
                        paragraph_lines = []
                        in_paragraph = False
                    html_parts.append(f'<h3>{line[4:]}</h3>')
                
                else:
                    paragraph_lines.append(line)
                    in_paragraph = True
            
            # Don't forget the last paragraph
            if paragraph_lines:
                html_parts.append(f'<p>{" ".join(paragraph_lines)}</p>')
        
        html_parts.extend(['</body>', '</html>'])
        return '\n'.join(html_parts)
    
    def _extract_headings_for_toc(self, content: str):
        """Extract headings from content for table of contents."""
        lines = content.split('\n')
        headings = []
        
        for i, line in enumerate(lines):
            line = line.strip()
            if line.startswith('# '):
                headings.append({
                    'level': 1,
                    'title': line[2:].strip(),
                    'line_number': i,
                    'anchor': f"chapter_{len([h for h in headings if h['level'] == 1]) + 1}"
                })
            elif line.startswith('## ') and not line.startswith('### '):
                headings.append({
                    'level': 2,
                    'title': line[3:].strip(),
                    'line_number': i,
                    'anchor': f"section_{len([h for h in headings if h['level'] == 2]) + 1}"
                })
            elif line.startswith('### '):
                headings.append({
                    'level': 3,
                    'title': line[4:].strip(),
                    'line_number': i,
                    'anchor': f"subsection_{len([h for h in headings if h['level'] == 3]) + 1}"
                })
        
        return headings
    
    def _create_professional_title_page(self, doc: Document, book, config: Dict[str, Any]):
        """Create a professional title page following publishing industry standards."""
        
        # Add vertical space before title (about 1/3 of page)
        for _ in range(8):
            doc.add_paragraph()
        
        # Main title - centered, large, bold
        title_p = doc.add_paragraph()
        title_run = title_p.add_run(book.title.upper())
        title_run.font.name = config['fonts']['title']
        title_run.font.size = Pt(config['font_sizes']['title'])
        title_run.font.bold = True
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.space_after = Pt(24)
        
        # Subtitle if exists (extract from title if has colon)
        if ':' in book.title:
            parts = book.title.split(':', 1)
            if len(parts) == 2:
                subtitle_p = doc.add_paragraph()
                subtitle_run = subtitle_p.add_run(parts[1].strip())
                subtitle_run.font.name = config['fonts']['subtitle']
                subtitle_run.font.size = Pt(config['font_sizes']['subtitle'])
                subtitle_run.font.bold = False
                subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                subtitle_p.space_after = Pt(36)
        
        # Add more vertical space
        for _ in range(6):
            doc.add_paragraph()
        
        # Author (Buko AI)
        author_p = doc.add_paragraph()
        author_run = author_p.add_run("Generado por Buko AI")
        author_run.font.name = config['fonts']['heading']
        author_run.font.size = Pt(config['font_sizes']['heading'])
        author_run.font.bold = False
        author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_p.space_after = Pt(48)
        
        # Add page break after title page
        doc.add_page_break()
        
        # Copyright/Information page
        self._create_copyright_page(doc, book, config)
    
    def _create_copyright_page(self, doc: Document, book, config: Dict[str, Any]):
        """Create a comprehensive professional copyright page."""
        
        # Copyright title
        copyright_title = doc.add_paragraph()
        title_run = copyright_title.add_run("Información de Publicación")
        title_run.font.name = config['fonts']['heading']
        title_run.font.size = Pt(config['font_sizes']['heading'])
        title_run.font.bold = True
        copyright_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        copyright_title.space_after = Pt(24)
        
        # Copyright notice
        copyright_notice = [
            f"© {datetime.now().year} {config['brand_elements']['publisher']}",
            "Todos los derechos reservados.",
            "",
            f"Título: {book.title}",
            f"Autor: Inteligencia Artificial - {config['brand_elements']['publisher']}",
            f"Editor: {config['brand_elements']['publisher']}",
            f"Primera edición digital: {datetime.now().strftime('%B de %Y')}",
            f"Género: {book.genre or 'Educativo'}",
            f"Idioma: {'Español' if book.language == 'es' else 'English' if book.language == 'en' else book.language or 'Español'}",
            f"Audiencia: {book.target_audience or 'General'}",
            "",
            "ISBN: No aplica (Contenido digital generado por IA)",
            f"Páginas: {book.page_count or 'Variable según formato'}",
            f"Palabras aproximadas: {book.final_words or 'Variable'}",
            "",
            "AVISO LEGAL:",
            "Este libro ha sido generado mediante inteligencia artificial avanzada.",
            "El contenido es con fines educativos e informativos únicamente.",
            "Aunque se ha puesto especial cuidado en la exactitud del contenido,",
            "el editor no garantiza la completitud o precisión de toda la información.",
            "",
            "Ninguna parte de esta publicación puede ser reproducida, almacenada",
            "en sistema de recuperación o transmitida en cualquier forma o por",
            "cualquier medio, electrónico, mecánico, fotocopia, grabación o",
            "cualquier otro, sin el permiso previo por escrito del editor.",
            "",
            f"Para más información, visite: {config['brand_elements']['website']}",
            f"Contacto: {config['brand_elements']['contact_email']}",
        ]
        
        for line in copyright_notice:
            if line:
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.name = config['fonts']['body_fallback']
                run.font.size = Pt(config['font_sizes']['metadata'])
                if line.startswith("©") or line == "AVISO LEGAL:":
                    run.font.bold = True
                p.space_after = Pt(4)
            else:
                doc.add_paragraph().space_after = Pt(8)
        
        doc.add_page_break()
        
        # Create additional professional pages if required
        if config.get('include_about_author'):
            self._create_about_author_page(doc, book, config)
        
        if config.get('include_disclaimer'):
            self._create_disclaimer_page(doc, book, config)
    
    def _create_about_author_page(self, doc: Document, book, config: Dict[str, Any]):
        """Create professional 'About the Author' page for Buko AI."""
        
        # Title
        title_p = doc.add_paragraph()
        title_run = title_p.add_run("Acerca de Buko AI")
        title_run.font.name = config['fonts']['heading']
        title_run.font.size = Pt(config['font_sizes']['title'])
        title_run.font.bold = True
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.space_after = Pt(24)
        
        # About content
        about_content = [
            "Buko AI es una plataforma revolucionaria de inteligencia artificial especializada",
            "en la creación de contenido educativo de alta calidad. Utilizando algoritmos",
            "avanzados de procesamiento de lenguaje natural y aprendizaje automático,",
            "Buko AI genera libros, guías y materiales educativos personalizados que se",
            "adaptan a las necesidades específicas de cada usuario.",
            "",
            "NUESTRA MISIÓN:",
            "Democratizar el acceso a contenido educativo de calidad mediante tecnología",
            "de inteligencia artificial, permitiendo que cualquier persona pueda acceder",
            "a materiales de aprendizaje personalizados, actualizados y relevantes.",
            "",
            "TECNOLOGÍA:",
            "• Procesamiento de lenguaje natural avanzado",
            "• Algoritmos de generación de contenido contextual",
            "• Personalización basada en perfil del usuario",
            "• Control de calidad automatizado",
            "• Formatos multiplataforma (PDF, EPUB, DOCX)",
            "",
            "CALIDAD Y PRECISIÓN:",
            "Aunque el contenido es generado por IA, cada libro pasa por procesos",
            "de verificación y estructuración para garantizar coherencia, relevancia",
            "y utilidad educativa. Nuestros algoritmos están continuamente aprendiendo",
            "y mejorando para ofrecer contenido cada vez más preciso y valioso.",
            "",
            f"Para obtener más libros personalizados, visite: {config['brand_elements']['website']}",
            f"Soporte y consultas: {config['brand_elements']['contact_email']}",
        ]
        
        for line in about_content:
            if line:
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.name = config['fonts']['body_fallback']
                run.font.size = Pt(config['font_sizes']['body'])
                
                if line.endswith(":") and line.isupper():
                    run.font.bold = True
                elif line.startswith("•"):
                    p.paragraph_format.left_indent = Cm(0.6)
                
                p.space_after = Pt(6)
            else:
                doc.add_paragraph().space_after = Pt(8)
        
        doc.add_page_break()
    
    def _create_disclaimer_page(self, doc: Document, book, config: Dict[str, Any]):
        """Create professional disclaimer and license page."""
        
        # Title
        title_p = doc.add_paragraph()
        title_run = title_p.add_run("Aviso Legal y Condiciones de Uso")
        title_run.font.name = config['fonts']['heading']
        title_run.font.size = Pt(config['font_sizes']['title'])
        title_run.font.bold = True
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.space_after = Pt(24)
        
        # Disclaimer content
        disclaimer_content = [
            "IMPORTANTE - LEA CUIDADOSAMENTE:",
            "",
            "1. NATURALEZA DEL CONTENIDO",
            "Este libro ha sido generado mediante inteligencia artificial. Aunque se han",
            "implementado controles de calidad, el contenido debe considerarse como material",
            "de referencia y apoyo educativo, no como autoridad definitiva en la materia.",
            "",
            "2. USO RESPONSABLE",  
            "El contenido está destinado únicamente para fines educativos, informativos",
            "y de entretenimiento. El usuario es responsable de verificar la información",
            "antes de tomar decisiones basadas en el contenido de este libro.",
            "",
            "3. LIMITACIÓN DE RESPONSABILIDAD",
            f"{config['brand_elements']['publisher']} no se hace responsable por:",
            "• Decisiones tomadas basándose en el contenido",
            "• Posibles imprecisiones o errores en la información",
            "• Resultados obtenidos al aplicar el contenido",
            "• Daños directos o indirectos derivados del uso",
            "",
            "4. DERECHOS DE AUTOR Y LICENCIA",
            "Este libro está protegido por derechos de autor. Se permite:",
            "• Uso personal y educativo",
            "• Cita con atribución apropiada",
            "",
            "NO se permite:",
            "• Reproducción comercial sin autorización",
            "• Modificación o alteración del contenido",
            "• Distribución masiva sin permiso escrito",
            "",
            "5. ACTUALIZACIONES Y MEJORAS",
            "Como el contenido es generado por IA en constante evolución, pueden",
            "existir versiones más actualizadas. Visite nuestro sitio web para",
            "acceder a las últimas versiones y mejoras.",
            "",
            "6. CONTACTO Y SOPORTE",
            f"Para consultas sobre licencias: {config['brand_elements']['contact_email']}",
            f"Sitio web oficial: {config['brand_elements']['website']}",
            "",
            f"Fecha de este aviso: {datetime.now().strftime('%d de %B de %Y')}",
        ]
        
        for line in disclaimer_content:
            if line:
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.name = config['fonts']['body_fallback']
                
                if line.isupper() or (line[0].isdigit() and line[1] == "."):
                    run.font.size = Pt(config['font_sizes']['subheading'])
                    run.font.bold = True
                    p.space_before = Pt(12)
                elif line.startswith("•") or line.startswith("NO"):
                    run.font.size = Pt(config['font_sizes']['metadata'])
                    p.paragraph_format.left_indent = Cm(0.8)
                else:
                    run.font.size = Pt(config['font_sizes']['metadata'])
                
                p.space_after = Pt(6)
            else:
                doc.add_paragraph().space_after = Pt(6)
        
        doc.add_page_break()
    
    def _create_professional_toc(self, doc: Document, headings: list, config: Dict[str, Any]):
        """Create a professional table of contents with real navigation links."""
        # TOC Title
        toc_title = doc.add_paragraph()
        title_run = toc_title.add_run("Tabla de Contenidos")
        title_run.font.name = config['fonts']['title']
        title_run.font.size = Pt(config['font_sizes']['title'])
        title_run.font.bold = True
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_title.space_after = Pt(24)
        
        # Add spacing before TOC entries
        doc.add_paragraph().space_after = Pt(12)
        
        # Generate TOC entries with navigation links
        for heading in headings:
            toc_entry = doc.add_paragraph()
            
            # Create hyperlink to bookmark
            self._add_hyperlink_to_bookmark(toc_entry, heading['title'], heading['bookmark_id'])
            
            # Style based on level
            if heading['level'] == 1:
                # Chapter entries - bold, no indent
                for run in toc_entry.runs:
                    run.font.name = config['fonts']['heading']
                    run.font.size = Pt(config['font_sizes']['toc'])
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 102, 204)  # Blue for links
                toc_entry.paragraph_format.left_indent = Cm(0)
                toc_entry.space_after = Pt(6)
                
            elif heading['level'] == 2:
                # Section entries - normal, small indent
                for run in toc_entry.runs:
                    run.font.name = config['fonts']['heading']
                    run.font.size = Pt(config['font_sizes']['toc'] - 1)
                    run.font.bold = False
                    run.font.color.rgb = RGBColor(0, 102, 204)  # Blue for links
                toc_entry.paragraph_format.left_indent = Cm(0.5)
                toc_entry.space_after = Pt(3)
                
            elif heading['level'] == 3:
                # Subsection entries - smaller, larger indent
                for run in toc_entry.runs:
                    run.font.name = config['fonts']['heading']
                    run.font.size = Pt(config['font_sizes']['toc'] - 2)
                    run.font.bold = False
                    run.font.color.rgb = RGBColor(0, 102, 204)  # Blue for links
                toc_entry.paragraph_format.left_indent = Cm(1.0)
                toc_entry.space_after = Pt(2)
        
        # Add page break after TOC
        doc.add_page_break()
    
    def _add_hyperlink_to_bookmark(self, paragraph, text: str, bookmark_name: str):
        """Add a hyperlink to a bookmark in the paragraph."""
        # Create hyperlink element
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('w:anchor'), bookmark_name)
        
        # Create run element
        run_element = OxmlElement('w:r')
        
        # Create text element
        text_element = OxmlElement('w:t')
        text_element.text = text
        
        # Assemble the hyperlink
        run_element.append(text_element)
        hyperlink.append(run_element)
        
        # Add to paragraph
        paragraph._element.append(hyperlink)
    
    def _add_bookmark(self, paragraph, bookmark_name: str):
        """Add a bookmark to a paragraph."""
        # Create bookmark start
        bookmark_start = OxmlElement('w:bookmarkStart')
        bookmark_start.set(qn('w:id'), str(hash(bookmark_name) % 1000000))
        bookmark_start.set(qn('w:name'), bookmark_name)
        
        # Create bookmark end
        bookmark_end = OxmlElement('w:bookmarkEnd')
        bookmark_end.set(qn('w:id'), str(hash(bookmark_name) % 1000000))
        
        # Insert at beginning and end of paragraph
        paragraph._element.insert(0, bookmark_start)
        paragraph._element.append(bookmark_end)
    
    def _process_markdown_formatting(self, paragraph, text: str):
        """Process markdown formatting like *italic* and **bold** with proper runs."""
        parts = []
        current_pos = 0
        
        # Find all markdown patterns
        import re
        patterns = [
            (r'\*\*(.+?)\*\*', 'bold'),
            (r'\*(.+?)\*', 'italic')
        ]
        
        # Find all matches and their positions
        matches = []
        for pattern, style in patterns:
            for match in re.finditer(pattern, text):
                matches.append({
                    'start': match.start(),
                    'end': match.end(),
                    'text': match.group(1),
                    'style': style
                })
        
        # Sort matches by position
        matches.sort(key=lambda x: x['start'])
        
        # Process text with formatting
        for match in matches:
            # Add text before the match
            if current_pos < match['start']:
                run = paragraph.add_run(text[current_pos:match['start']])
            
            # Add formatted text
            run = paragraph.add_run(match['text'])
            if match['style'] == 'bold':
                run.font.bold = True
            elif match['style'] == 'italic':
                run.font.italic = True
            
            current_pos = match['end']
        
        # Add remaining text
        if current_pos < len(text):
            paragraph.add_run(text[current_pos:])
        
        # If no formatting found, add as simple text
        if not matches:
            paragraph.add_run(text)
    
    def _extract_headings_for_toc_with_bookmarks(self, content: str) -> List[Dict[str, Any]]:
        """Extract headings with unique bookmark IDs for navigation."""
        headings = []
        lines = content.split('\n')
        
        chapter_count = 0
        section_count = 0
        subsection_count = 0
        
        for line in lines:
            line = line.strip()
            
            if line.startswith('# ') and not line.startswith('## '):
                chapter_count += 1
                title = line[2:].strip()
                bookmark_id = f"chapter_{chapter_count}_{title[:20].replace(' ', '_')}"
                headings.append({
                    'level': 1,
                    'title': title,
                    'bookmark_id': bookmark_id
                })
            
            elif line.startswith('## ') and not line.startswith('### '):
                section_count += 1
                title = line[3:].strip()
                bookmark_id = f"section_{section_count}_{title[:20].replace(' ', '_')}"
                headings.append({
                    'level': 2,
                    'title': title,
                    'bookmark_id': bookmark_id
                })
            
            elif line.startswith('### '):
                subsection_count += 1
                title = line[4:].strip()
                bookmark_id = f"subsection_{subsection_count}_{title[:20].replace(' ', '_')}"
                headings.append({
                    'level': 3,
                    'title': title,
                    'bookmark_id': bookmark_id
                })
        
        return headings
    
    def _add_professional_paragraph_with_formatting(self, doc: Document, text: str, config: Dict[str, Any], is_first_in_section=False):
        """Add a professionally formatted paragraph with markdown formatting support."""
        p = doc.add_paragraph()
        
        # Process markdown formatting with proper runs
        self._process_markdown_formatting(p, text.strip())
        
        # Set paragraph style
        for run in p.runs:
            run.font.name = config['fonts']['body']
            run.font.size = Pt(config['font_sizes']['body'])
        
        # Professional paragraph formatting
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(config['paragraph_spacing']['after'])
        p.paragraph_format.first_line_indent = Pt(config['paragraph_spacing']['first_line_indent']) if not is_first_in_section else Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = config['paragraph_spacing']['line_spacing']
    
    def _process_content_for_docx(self, doc: Document, content: str, config: Dict[str, Any]):
        """Process book content for DOCX document with professional formatting."""
        if not content:
            return
        
        # Extract headings for TOC  
        headings = self._extract_headings_for_toc(content)
        
        # Create professional table of contents
        if config.get('toc_required', True):
            self._create_professional_toc(doc, headings, config)
        
        lines = content.split('\n')
        paragraph_lines = []
        current_chapter = 0
        current_section = 0
        is_first_paragraph_in_section = True
        
        for line in lines:
            line = line.strip()
            
            if not line:
                # Empty line - end current paragraph if we have content
                if paragraph_lines:
                    full_text = ' '.join(paragraph_lines)
                    self._add_professional_paragraph_with_formatting(doc, full_text, config, is_first_paragraph_in_section)
                    paragraph_lines = []
                    is_first_paragraph_in_section = False
                # Don't add empty paragraphs - use spacing instead
            
            elif line.startswith('# '):
                # Main chapter title
                if paragraph_lines:
                    full_text = ' '.join(paragraph_lines)
                    self._add_professional_paragraph_with_formatting(doc, full_text, config)
                    paragraph_lines = []
                
                current_chapter += 1
                if config.get('chapter_breaks', True) and current_chapter > 1:
                    doc.add_page_break()
                
                # Chapter title with professional formatting
                chapter_p = doc.add_paragraph()
                chapter_run = chapter_p.add_run(line[2:].strip())
                chapter_run.font.name = config['fonts']['chapter']
                chapter_run.font.size = Pt(config['font_sizes']['chapter'])
                chapter_run.font.bold = True
                chapter_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                chapter_p.paragraph_format.space_before = Pt(config['heading_spacing']['chapter_before'])
                chapter_p.paragraph_format.space_after = Pt(config['heading_spacing']['chapter_after'])
                
                is_first_paragraph_in_section = True
            
            elif line.startswith('## ') and not line.startswith('### '):
                # Section heading
                if paragraph_lines:
                    full_text = ' '.join(paragraph_lines)
                    self._add_professional_paragraph_with_formatting(doc, full_text, config)
                    paragraph_lines = []
                
                current_section += 1
                
                # Professional section heading
                section_p = doc.add_paragraph()
                section_run = section_p.add_run(line[3:].strip())
                section_run.font.name = config['fonts']['heading']
                section_run.font.size = Pt(config['font_sizes']['heading'])
                section_run.font.bold = True
                section_p.paragraph_format.space_before = Pt(config['heading_spacing']['section_before'])
                section_p.paragraph_format.space_after = Pt(config['heading_spacing']['section_after'])
                
                is_first_paragraph_in_section = True
            
            elif line.startswith('### '):
                # Subsection heading
                if paragraph_lines:
                    full_text = ' '.join(paragraph_lines)
                    self._add_professional_paragraph_with_formatting(doc, full_text, config)
                    paragraph_lines = []
                
                # Professional subsection heading
                subsection_p = doc.add_paragraph()
                subsection_run = subsection_p.add_run(line[4:].strip())
                subsection_run.font.name = config['fonts']['heading']
                subsection_run.font.size = Pt(config['font_sizes']['subheading'])
                subsection_run.font.bold = True
                subsection_p.paragraph_format.space_before = Pt(config['heading_spacing']['subsection_before'])
                subsection_p.paragraph_format.space_after = Pt(config['heading_spacing']['subsection_after'])
                
                is_first_paragraph_in_section = True
            
            elif line.startswith('- ') or line.startswith('* '):
                # Bullet list item
                if paragraph_lines:
                    full_text = ' '.join(paragraph_lines)
                    self._add_professional_paragraph_with_formatting(doc, full_text, config)
                    paragraph_lines = []
                
                # Professional bullet point with markdown formatting
                bullet_p = doc.add_paragraph()
                bullet_text = f"• {line[2:].strip()}"
                self._process_markdown_formatting(bullet_p, bullet_text)
                
                # Style bullet points
                for run in bullet_p.runs:
                    run.font.name = config['fonts']['body']
                    run.font.size = Pt(config['font_sizes']['body'])
                
                bullet_p.paragraph_format.left_indent = Cm(0.6)
                bullet_p.paragraph_format.space_after = Pt(3)
                
                is_first_paragraph_in_section = False
            
            else:
                # Regular text - accumulate for paragraph
                paragraph_lines.append(line)
        
        # Don't forget the last paragraph
        if paragraph_lines:
            full_text = ' '.join(paragraph_lines)
            self._add_professional_paragraph_with_formatting(doc, full_text, config)
    
    def _process_professional_content_for_pdf(self, book_structure: 'BookStructure', styles: Dict[str, Any]) -> List:
        """Process professionally formatted BookStructure content for PDF."""
        story = []
        
        if not book_structure or not book_structure.elements:
            return story
        
        for element in book_structure.elements:
            element_type = element.type.value if hasattr(element.type, 'value') else str(element.type)
            content = element.content or ""
            
            if element_type == 'book-title':
                story.append(Paragraph(content, styles.get('BookTitle', styles['Title'])))
                story.append(Spacer(1, 0.5*inch))
                
            elif element_type in ['chapter', 'chapter-title']:
                if content:
                    story.append(Paragraph(content, styles.get('ChapterHeading', styles['Heading1'])))
                    
            elif element_type in ['section', 'section-title']:
                if content:
                    story.append(Paragraph(content, styles.get('SectionHeading', styles['Heading2'])))
                    
            elif element_type in ['subsection', 'subsection-title']:
                if content:
                    story.append(Paragraph(content, styles.get('SectionHeading', styles['Heading3'])))
                    
            elif element_type == 'paragraph':
                if content.strip():
                    # Convert HTML to plain text for PDF
                    if BS4_AVAILABLE:
                        clean_content = BeautifulSoup(content, 'html.parser').get_text()
                    else:
                        # Fallback: simple HTML tag removal
                        import re
                        clean_content = re.sub(r'<[^>]+>', '', content)
                    story.append(Paragraph(clean_content, styles.get('BookBody', styles['Normal'])))
                    
            elif element_type == 'expression':
                if content.strip():
                    # Handle numbered expressions with special formatting
                    if BS4_AVAILABLE:
                        clean_content = BeautifulSoup(content, 'html.parser').get_text()
                    else:
                        import re
                        clean_content = re.sub(r'<[^>]+>', '', content)
                    # Add some visual distinction for expressions
                    story.append(Spacer(1, 6))
                    story.append(Paragraph(f"• {clean_content}", styles.get('BookBody', styles['Normal'])))
                    story.append(Spacer(1, 6))
                    
            elif element_type in ['list-item', 'item']:
                if content.strip():
                    if BS4_AVAILABLE:
                        clean_content = BeautifulSoup(content, 'html.parser').get_text()
                    else:
                        import re
                        clean_content = re.sub(r'<[^>]+>', '', content)
                    story.append(Paragraph(f"• {clean_content}", styles.get('BookBody', styles['Normal'])))
                    
            # Add spacing between elements
            if element_type in ['paragraph', 'expression', 'list-item']:
                story.append(Spacer(1, 6))
        
        return story
    
    def _process_professional_content_for_epub(self, book_structure: 'BookStructure', config: Dict[str, Any]) -> List[Tuple[str, str]]:
        """Process professionally formatted BookStructure content for EPUB."""
        chapters = []
        
        if not book_structure or not book_structure.elements:
            return chapters
        
        # Group elements by chapters
        current_chapter_title = "Introducción"
        current_chapter_elements = []
        
        for element in book_structure.elements:
            element_type = element.type.value if hasattr(element.type, 'value') else str(element.type)
            content = element.content or ""
            
            if element_type in ['chapter', 'chapter-title'] and content.strip():
                # Save previous chapter if exists
                if current_chapter_elements:
                    html_content = self._format_professional_epub_chapter(current_chapter_title, current_chapter_elements)
                    chapters.append((current_chapter_title, html_content))
                
                # Start new chapter
                current_chapter_title = content.strip()
                current_chapter_elements = []
            else:
                current_chapter_elements.append(element)
        
        # Don't forget the last chapter
        if current_chapter_elements:
            html_content = self._format_professional_epub_chapter(current_chapter_title, current_chapter_elements)
            chapters.append((current_chapter_title, html_content))
        
        return chapters
    
    def _format_professional_epub_chapter(self, title: str, elements: List) -> str:
        """Format a chapter from professional elements for EPUB."""
        html_parts = [
            '<?xml version="1.0" encoding="UTF-8"?>',
            '<!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.1//EN" "http://www.w3.org/TR/xhtml11/DTD/xhtml11.dtd">',
            '<html xmlns="http://www.w3.org/1999/xhtml">',
            '<head>',
            f'<title>{title}</title>',
            '<link rel="stylesheet" type="text/css" href="style/main.css"/>',
            '</head>',
            '<body>',
            f'<h2>{title}</h2>'
        ]
        
        for element in elements:
            element_type = element.type.value if hasattr(element.type, 'value') else str(element.type)
            content = element.content or ""
            
            if not content.strip():
                continue
                
            if element_type in ['section', 'section-title']:
                html_parts.append(f'<h3>{content}</h3>')
            elif element_type in ['subsection', 'subsection-title']:
                html_parts.append(f'<h4>{content}</h4>')
            elif element_type == 'paragraph':
                # Content is already HTML from professional formatting
                html_parts.append(f'<p>{content}</p>')
            elif element_type == 'expression':
                # Handle expressions with special formatting
                html_parts.append(f'<div class="expression"><p>{content}</p></div>')
            elif element_type in ['list-item', 'item']:
                html_parts.append(f'<p>• {content}</p>')
            else:
                # Default paragraph treatment
                html_parts.append(f'<p>{content}</p>')
        
        html_parts.extend(['</body>', '</html>'])
        return '\n'.join(html_parts)
    
    def _process_professional_content_for_docx(self, doc: Document, book_structure: 'BookStructure', config: Dict[str, Any]):
        """Process professionally formatted BookStructure content for DOCX."""
        if not book_structure or not book_structure.elements:
            return
        
        # Create table of contents if elements are available
        if config.get('toc_required', True) and book_structure.toc:
            self._create_professional_toc_from_structure(doc, book_structure, config)
        
        for element in book_structure.elements:
            element_type = element.type.value if hasattr(element.type, 'value') else str(element.type)
            content = element.content or ""
            
            if not content.strip():
                continue
            
            if element_type == 'book-title':
                # Main book title
                title_p = doc.add_paragraph()
                title_run = title_p.add_run(content.upper())
                title_run.font.name = config['fonts']['title']
                title_run.font.size = Pt(config['font_sizes']['title'])
                title_run.font.bold = True
                title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_p.space_after = Pt(36)
                
            elif element_type in ['chapter', 'chapter-title']:
                # Chapter heading with page break
                if config.get('chapter_breaks', True):
                    doc.add_page_break()
                
                chapter_p = doc.add_paragraph()
                chapter_run = chapter_p.add_run(content)
                chapter_run.font.name = config['fonts']['chapter']
                chapter_run.font.size = Pt(config['font_sizes']['chapter'])
                chapter_run.font.bold = True
                chapter_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                chapter_p.paragraph_format.space_before = Pt(config['heading_spacing']['chapter_before'])
                chapter_p.paragraph_format.space_after = Pt(config['heading_spacing']['chapter_after'])
                
            elif element_type in ['section', 'section-title']:
                # Section heading
                section_p = doc.add_paragraph()
                section_run = section_p.add_run(content)
                section_run.font.name = config['fonts']['heading']
                section_run.font.size = Pt(config['font_sizes']['heading'])
                section_run.font.bold = True
                section_p.paragraph_format.space_before = Pt(config['heading_spacing']['section_before'])
                section_p.paragraph_format.space_after = Pt(config['heading_spacing']['section_after'])
                
            elif element_type in ['subsection', 'subsection-title']:
                # Subsection heading
                subsection_p = doc.add_paragraph()
                subsection_run = subsection_p.add_run(content)
                subsection_run.font.name = config['fonts']['heading']
                subsection_run.font.size = Pt(config['font_sizes']['subheading'])
                subsection_run.font.bold = True
                subsection_p.paragraph_format.space_before = Pt(config['heading_spacing']['subsection_before'])
                subsection_p.paragraph_format.space_after = Pt(config['heading_spacing']['subsection_after'])
                
            elif element_type == 'paragraph':
                # Regular paragraph with HTML content support
                if BS4_AVAILABLE:
                    soup = BeautifulSoup(content, 'html.parser')
                    clean_text = soup.get_text()
                else:
                    import re
                    clean_text = re.sub(r'<[^>]+>', '', content)
                
                p = doc.add_paragraph()
                self._process_markdown_formatting(p, clean_text)
                
                # Style the paragraph
                for run in p.runs:
                    run.font.name = config['fonts']['body']
                    run.font.size = Pt(config['font_sizes']['body'])
                
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_after = Pt(config['paragraph_spacing']['after'])
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                p.paragraph_format.line_spacing = config['line_spacing']
                
            elif element_type == 'expression':
                # Special formatting for expressions
                if BS4_AVAILABLE:
                    soup = BeautifulSoup(content, 'html.parser')
                    clean_text = soup.get_text()
                else:
                    import re
                    clean_text = re.sub(r'<[^>]+>', '', content)
                
                expr_p = doc.add_paragraph()
                expr_run = expr_p.add_run(f"• {clean_text}")
                expr_run.font.name = config['fonts']['body']
                expr_run.font.size = Pt(config['font_sizes']['body'])
                expr_run.font.bold = True
                
                expr_p.paragraph_format.left_indent = Cm(0.8)
                expr_p.paragraph_format.space_after = Pt(8)
                
            elif element_type in ['list-item', 'item']:
                # List items
                if BS4_AVAILABLE:
                    soup = BeautifulSoup(content, 'html.parser')
                    clean_text = soup.get_text()
                else:
                    import re
                    clean_text = re.sub(r'<[^>]+>', '', content)
                
                item_p = doc.add_paragraph()
                item_run = item_p.add_run(f"• {clean_text}")
                item_run.font.name = config['fonts']['body']
                item_run.font.size = Pt(config['font_sizes']['body'])
                
                item_p.paragraph_format.left_indent = Cm(0.6)
                item_p.paragraph_format.space_after = Pt(4)
    
    def _create_professional_toc_from_structure(self, doc: Document, book_structure: 'BookStructure', config: Dict[str, Any]):
        """Create TOC from professional book structure."""
        if not book_structure.toc:
            return
            
        # TOC Title
        toc_title = doc.add_paragraph()
        title_run = toc_title.add_run("Tabla de Contenidos")
        title_run.font.name = config['fonts']['title']
        title_run.font.size = Pt(config['font_sizes']['title'])
        title_run.font.bold = True
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_title.space_after = Pt(24)
        
        # Add TOC entries
        for toc_entry in book_structure.toc:
            if not toc_entry.get('title'):
                continue
                
            entry_p = doc.add_paragraph()
            entry_run = entry_p.add_run(toc_entry['title'])
            entry_run.font.name = config['fonts']['heading']
            entry_run.font.size = Pt(config['font_sizes']['toc'])
            
            # Indent based on level
            level = toc_entry.get('level', 1)
            if level == 1:
                entry_run.font.bold = True
                entry_p.paragraph_format.left_indent = Cm(0)
            elif level == 2:
                entry_p.paragraph_format.left_indent = Cm(0.5)
            else:
                entry_p.paragraph_format.left_indent = Cm(1.0)
            
            entry_p.space_after = Pt(6)
        
        # Add page break after TOC
        doc.add_page_break()