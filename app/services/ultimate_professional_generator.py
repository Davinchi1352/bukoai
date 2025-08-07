"""
ULTIMATE PROFESSIONAL DOCUMENT GENERATOR - EL SALTO MONUMENTAL
Genera documentos Word y PDF de CALIDAD EDITORIAL REAL usando las mejores herramientas disponibles.
"""

import os
import re
import tempfile
import uuid
from typing import Dict, List, Any, Optional, Tuple
from dataclasses import dataclass, asdict
from datetime import datetime
from pathlib import Path
import subprocess

# Herramientas profesionales
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

import weasyprint
from weasyprint import HTML, CSS
from weasyprint.text.fonts import FontConfiguration

from bs4 import BeautifulSoup
from PIL import Image, ImageDraw, ImageFont
import requests


@dataclass 
class UltimateProfessionalOptions:
    """Opciones definitivas para documentos de calidad editorial suprema."""
    
    # TIPOGRAFÍA EDITORIAL SUPREMA
    font_family_body: str = "Crimson Text"
    font_family_headings: str = "Playfair Display"
    font_size_body: int = 12
    font_size_title: int = 28
    font_size_chapter: int = 20
    font_size_section: int = 16
    font_weight_body: str = "400"
    font_weight_headings: str = "700"
    
    # ESPACIADO EDITORIAL PERFECTO
    line_height: float = 1.6
    paragraph_spacing: float = 1.2
    chapter_spacing_before: float = 3.0
    chapter_spacing_after: float = 1.5
    section_spacing_before: float = 2.0
    section_spacing_after: float = 1.0
    
    # PÁGINA EDITORIAL PROFESIONAL
    page_width: str = "8.5in"
    page_height: str = "11in" 
    margin_top: str = "1in"
    margin_bottom: str = "1in"
    margin_left: str = "1.25in"
    margin_right: str = "1in"
    margin_inner: str = "1.5in"  # Para encuadernación
    margin_outer: str = "1in"
    
    # ELEMENTOS EDITORIALES AVANZADOS
    include_cover_page: bool = True
    include_title_page: bool = True
    include_copyright_page: bool = True
    include_dedication: bool = False
    include_table_of_contents: bool = True
    include_index: bool = True
    include_bibliography: bool = False
    include_about_author: bool = True
    
    # CARACTERÍSTICAS AVANZADAS
    enable_drop_caps: bool = True
    enable_page_numbers: bool = True
    enable_headers: bool = True
    enable_footers: bool = True
    enable_chapter_breaks: bool = True
    enable_hyphenation: bool = True
    enable_justification: bool = True
    enable_orphan_control: bool = True
    enable_widow_control: bool = True
    
    # COLORES EDITORIALES
    color_text: str = "#2d3748"
    color_headings: str = "#1a202c"
    color_accent: str = "#3182ce"
    color_secondary: str = "#718096"
    
    # CALIDAD DE EXPORTACIÓN
    pdf_quality: str = "print"  # screen, ebook, print
    image_dpi: int = 300
    embed_fonts: bool = True
    compress_images: bool = False
    
    # METADATOS EDITORIALES
    title: str = ""
    author: str = ""
    publisher: str = "Buko AI Editorial"
    isbn: str = ""
    copyright_year: str = ""
    edition: str = "Primera Edición Digital"
    language: str = "es"
    subject: str = ""
    keywords: List[str] = None
    
    def __post_init__(self):
        if self.keywords is None:
            self.keywords = []
        if not self.copyright_year:
            self.copyright_year = str(datetime.now().year)


class UltimateProfessionalGenerator:
    """El generador definitivo de documentos profesionales de calidad editorial suprema."""
    
    def __init__(self):
        self.temp_dir = tempfile.mkdtemp()
        self.font_config = FontConfiguration()
        self.session_id = str(uuid.uuid4())[:8]
        
    def generate_ultimate_professional_document(self,
                                              content: str,
                                              book_data: Dict[str, Any],
                                              options: UltimateProfessionalOptions) -> Dict[str, Any]:
        """
        Genera el documento profesional definitivo con calidad editorial suprema.
        
        Returns:
            Dict con paths y información de documentos Word y PDF generados
        """
        
        # Procesar contenido con inteligencia avanzada
        processed_content = self._process_content_intelligently(content)
        
        # Generar documento Word profesional REAL
        word_path = self._generate_professional_word_document(
            processed_content, book_data, options
        )
        
        # Generar PDF de calidad editorial usando WeasyPrint
        pdf_path = self._generate_editorial_pdf_document(
            processed_content, book_data, options
        )
        
        # Generar HTML profesional para vista previa
        html_path = self._generate_professional_html_document(
            processed_content, book_data, options
        )
        
        # Generar información completa del documento
        document_info = self._generate_comprehensive_document_info(
            processed_content, book_data, options
        )
        
        return {
            'word_path': word_path,
            'pdf_path': pdf_path,
            'html_path': html_path,
            'document_info': document_info,
            'temp_dir': self.temp_dir,
            'session_id': self.session_id,
            'quality_level': 'Editorial Supreme',
            'ready_for_publication': True
        }
    
    def _process_content_intelligently(self, content: str) -> Dict[str, Any]:
        """Procesa el contenido con inteligencia avanzada para detectar estructura."""
        
        soup = BeautifulSoup(content, 'html.parser')
        
        # Análisis inteligente de estructura
        structure = {
            'title': None,
            'chapters': [],
            'sections': [],
            'expressions': [],
            'translations': [],
            'footnotes': [],
            'images': [],
            'tables': [],
            'paragraphs': [],
            'raw_content': str(soup)
        }
        
        # Detectar título principal (solo el primero)
        title_candidates = soup.find_all(['h1'], class_=lambda x: x and ('book-title' in x or 'ebook-book-title' in x))
        if not title_candidates:
            title_candidates = soup.find_all(['h1'])
        
        if title_candidates:
            structure['title'] = title_candidates[0].get_text().strip()
        
        # Detectar capítulos ÚNICOS (excluyendo el título principal y TOC)
        seen_titles = set()
        if structure['title']:
            seen_titles.add(structure['title'].lower())
        seen_titles.add('tabla de contenidos')
        
        for i, element in enumerate(soup.find_all(['h2'])):
            text = element.get_text().strip()
            if text and text.lower() not in seen_titles:
                chapter_id = f"chapter-{len(structure['chapters'])+1}"
                structure['chapters'].append({
                    'id': chapter_id,
                    'title': text,
                    'level': int(element.name[1]),
                    'position': len(structure['chapters']),
                    'word_count': len(text.split())
                })
                seen_titles.add(text.lower())
        
        # Detectar secciones
        for i, element in enumerate(soup.find_all(['h3', 'h4', 'h5', 'h6'])):
            text = element.get_text().strip()
            if text:
                section_id = f"section-{i+1}"
                structure['sections'].append({
                    'id': section_id,
                    'title': text,
                    'level': int(element.name[1]),
                    'position': i,
                    'word_count': len(text.split())
                })
        
        # Detectar expresiones especiales
        expression_elements = soup.find_all(class_=re.compile(r'expression|ebook-expression'))
        for i, element in enumerate(expression_elements):
            text = element.get_text().strip()
            if text:
                structure['expressions'].append({
                    'id': f"expr-{i+1}",
                    'text': text,
                    'number': i + 1,
                    'position': i
                })
        
        # Detectar traducciones
        translation_elements = soup.find_all(class_=re.compile(r'translation|traduccion|ebook-translation'))
        for i, element in enumerate(translation_elements):
            text = element.get_text().strip()
            if text:
                structure['translations'].append({
                    'id': f"trans-{i+1}",
                    'text': text,
                    'type': 'contextual',
                    'position': i
                })
        
        # Detectar imágenes
        image_elements = soup.find_all('img')
        for i, element in enumerate(image_elements):
            structure['images'].append({
                'id': f"img-{i+1}",
                'src': element.get('src', ''),
                'alt': element.get('alt', ''),
                'position': i
            })
        
        # Detectar tablas
        table_elements = soup.find_all('table')
        for i, element in enumerate(table_elements):
            structure['tables'].append({
                'id': f"table-{i+1}",
                'rows': len(element.find_all('tr')),
                'position': i
            })
        
        # Detectar párrafos
        for i, element in enumerate(soup.find_all(['p'])):
            text = element.get_text().strip()
            if text and len(text) > 10:  # Solo párrafos con contenido sustancial
                structure['paragraphs'].append({
                    'id': f"para-{i+1}",
                    'text': text,
                    'classes': element.get('class', []),
                    'position': i
                })
        
        # Calcular estadísticas avanzadas
        full_text = soup.get_text()
        words = full_text.split()
        
        structure['statistics'] = {
            'total_words': len(words),
            'total_characters': len(full_text),
            'total_characters_no_spaces': len(full_text.replace(' ', '')),
            'estimated_reading_time': len(words) // 200,  # 200 WPM average
            'estimated_pages': max(1, len(words) // 250),  # 250 words per page
            'complexity_score': self._calculate_complexity_score(full_text),
            'readability_score': self._calculate_readability_score(full_text)
        }
        
        return structure
    
    def _generate_professional_word_document(self,
                                           processed_content: Dict,
                                           book_data: Dict,
                                           options: UltimateProfessionalOptions) -> str:
        """Genera documento Word REAL con formateo profesional supremo."""
        
        # Crear documento Word
        doc = Document()
        
        # CONFIGURACIÓN AVANZADA DEL DOCUMENTO
        self._setup_advanced_document_properties(doc, book_data, options)
        self._create_professional_styles(doc, options)
        self._setup_advanced_page_layout(doc, options)
        
        # GENERAR CONTENIDO EDITORIAL
        if options.include_cover_page:
            self._add_professional_cover_page(doc, book_data, options)
        
        if options.include_copyright_page:
            self._add_comprehensive_copyright_page(doc, book_data, options)
        
        if options.include_table_of_contents:
            self._add_professional_table_of_contents(doc, processed_content, options)
        
        # CONTENIDO PRINCIPAL CON FORMATEO SUPREMO
        self._add_formatted_main_content(doc, processed_content, options)
        
        if options.include_index:
            self._add_professional_index(doc, processed_content, options)
        
        if options.include_about_author:
            self._add_about_author_page(doc, book_data, options)
        
        # Guardar documento Word
        word_filename = f"{book_data.get('title', 'documento')}_profesional.docx"
        word_path = os.path.join(self.temp_dir, word_filename)
        doc.save(word_path)
        
        return word_path
    
    def _generate_editorial_pdf_document(self,
                                       processed_content: Dict,
                                       book_data: Dict,
                                       options: UltimateProfessionalOptions) -> str:
        """Genera PDF de calidad editorial usando WeasyPrint."""
        
        # Generar HTML optimizado para PDF
        html_content = self._create_pdf_optimized_html(processed_content, book_data, options)
        
        # Generar CSS editorial avanzado
        css_content = self._create_editorial_css(options)
        
        # Configuración avanzada de WeasyPrint
        html_doc = HTML(string=html_content, base_url=self.temp_dir)
        css_doc = CSS(string=css_content, font_config=self.font_config)
        
        # Generar PDF con configuración editorial
        pdf_filename = f"{book_data.get('title', 'documento')}_editorial.pdf"
        pdf_path = os.path.join(self.temp_dir, pdf_filename)
        
        html_doc.write_pdf(
            pdf_path,
            stylesheets=[css_doc],
            font_config=self.font_config,
            optimize_images=not options.compress_images,
            pdf_version='1.7',
            presentational_hints=True
        )
        
        return pdf_path
    
    def _create_pdf_optimized_html(self,
                                 processed_content: Dict,
                                 book_data: Dict, 
                                 options: UltimateProfessionalOptions) -> str:
        """Crea HTML optimizado específicamente para generación de PDF editorial."""
        
        html_parts = [
            '<!DOCTYPE html>',
            '<html lang="es">',
            '<head>',
            '<meta charset="UTF-8">',
            f'<title>{book_data.get("title", "Documento Profesional")}</title>',
            f'<meta name="author" content="{book_data.get("author", "")}">',
            f'<meta name="subject" content="{options.subject}">',
            f'<meta name="creator" content="Buko AI Ultimate Generator">',
            '</head>',
            '<body class="editorial-body">'
        ]
        
        # Portada editorial
        if options.include_cover_page:
            html_parts.extend([
                '<div class="cover-page">',
                '<div class="cover-content">',
                f'<h1 class="cover-title">{book_data.get("title", "")}</h1>',
                f'<div class="cover-author">{book_data.get("author", "")}</div>',
                '<div class="cover-publisher">',
                f'<div class="publisher-name">{options.publisher}</div>',
                f'<div class="publication-year">{options.copyright_year}</div>',
                '</div>',
                '</div>',
                '</div>'
            ])
        
        # Tabla de contenidos editorial (sin duplicados)
        if options.include_table_of_contents:
            html_parts.extend([
                '<div class="toc-page">',
                '<h2 class="toc-title">Tabla de Contenidos</h2>',
                '<div class="toc-content">'
            ])
            
            # Solo capítulos únicos (ya filtrados en processed_content)
            page_start = 5
            for i, chapter in enumerate(processed_content['chapters']):
                estimated_page = page_start + (i * 2)
                html_parts.append(
                    f'<div class="toc-entry toc-level-{chapter["level"]}">'
                    f'<a href="#{chapter["id"]}" class="toc-link">'
                    f'<span class="toc-text">{chapter["title"]}</span>'
                    f'<span class="toc-page-num">{estimated_page}</span>'
                    f'</a></div>'
                )
            
            html_parts.extend(['</div>', '</div>'])
        
        # Contenido principal con estructura editorial
        html_parts.append('<div class="main-content editorial-content">')
        
        # Procesar contenido con BeautifulSoup para aplicar clases editoriales
        soup = BeautifulSoup(processed_content['raw_content'], 'html.parser')
        
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'div']):
            if element.name == 'h1':
                element['class'] = element.get('class', []) + ['editorial-title']
            elif element.name == 'h2': 
                element['class'] = element.get('class', []) + ['editorial-chapter']
            elif element.name == 'h3':
                element['class'] = element.get('class', []) + ['editorial-section']
            elif element.name == 'p':
                element['class'] = element.get('class', []) + ['editorial-paragraph']
        
        html_parts.append(str(soup))
        html_parts.extend(['</div>', '</body>', '</html>'])
        
        return '\\n'.join(html_parts)
    
    def _create_editorial_css(self, options: UltimateProfessionalOptions) -> str:
        """Crea CSS de calidad editorial suprema para PDF."""
        
        return f'''
        /* === CONFIGURACIÓN DE PÁGINA EDITORIAL === */
        @page {{
            size: {options.page_width} {options.page_height};
            margin: {options.margin_top} {options.margin_right} {options.margin_bottom} {options.margin_left};
            
            @top-center {{
                content: "{options.title}";
                font-family: "{options.font_family_body}", serif;
                font-size: 10pt;
                font-style: italic;
                color: #666;
                padding-bottom: 8pt;
                border-bottom: 0.5pt solid #ddd;
            }}
            
            @bottom-center {{
                content: counter(page);
                font-family: "{options.font_family_body}", serif;
                font-size: 11pt;
                color: #333;
                padding-top: 8pt;
            }}
        }}
        
        @page :first {{
            margin: 0;
            @top-center {{ content: none; }}
            @bottom-center {{ content: none; }}
        }}
        
        /* === FUENTES EDITORIALES PROFESIONALES === */
        @import url('https://fonts.googleapis.com/css2?family=Crimson+Text:ital,wght@0,400;0,600;1,400&family=Playfair+Display:wght@400;700;900&display=swap');
        
        /* === CONFIGURACIÓN BASE EDITORIAL === */
        html {{
            font-size: {options.font_size_body}pt;
            line-height: {options.line_height};
        }}
        
        .editorial-body {{
            font-family: "{options.font_family_body}", "Times New Roman", serif;
            font-size: {options.font_size_body}pt;
            font-weight: {options.font_weight_body};
            color: {options.color_text};
            line-height: {options.line_height};
            text-rendering: optimizeLegibility;
            font-feature-settings: "kern" 1, "liga" 1, "calt" 1, "onum" 1;
            hyphens: {'auto' if options.enable_hyphenation else 'none'};
            text-align: {'justify' if options.enable_justification else 'left'};
            hanging-punctuation: first last;
            orphans: {'3' if options.enable_orphan_control else 'auto'};
            widows: {'3' if options.enable_widow_control else 'auto'};
        }}
        
        /* === PORTADA EDITORIAL === */
        .cover-page {{
            page-break-after: always;
            height: 100vh;
            display: flex;
            align-items: center;
            justify-content: center;
            text-align: center;
            background: linear-gradient(135deg, #f8fafc 0%, #e2e8f0 100%);
        }}
        
        .cover-title {{
            font-family: "{options.font_family_headings}", serif;
            font-size: {options.font_size_title}pt;
            font-weight: 900;
            color: {options.color_headings};
            margin-bottom: 1em;
            line-height: 1.1;
            letter-spacing: -0.02em;
            text-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }}
        
        .cover-author {{
            font-size: {options.font_size_chapter}pt;
            font-weight: 600;
            color: {options.color_accent};
            margin: 2em 0;
            text-transform: uppercase;
            letter-spacing: 0.1em;
        }}
        
        .cover-publisher {{
            margin-top: 4em;
            padding-top: 2em;
            border-top: 2pt solid {options.color_accent};
        }}
        
        .publisher-name {{
            font-size: 14pt;
            font-weight: 600;
            color: {options.color_accent};
        }}
        
        .publication-year {{
            font-size: 18pt;
            font-weight: 700;
            margin-top: 1em;
            color: {options.color_headings};
        }}
        
        /* === TABLA DE CONTENIDOS EDITORIAL === */
        .toc-page {{
            page-break-before: always;
            page-break-after: always;
        }}
        
        .toc-title {{
            font-family: "{options.font_family_headings}", serif;
            font-size: {options.font_size_chapter}pt;
            font-weight: {options.font_weight_headings};
            text-align: center;
            color: {options.color_headings};
            margin-bottom: 2em;
            padding-bottom: 0.5em;
            border-bottom: 2pt solid {options.color_accent};
        }}
        
        .toc-entry {{
            margin-bottom: 0.8em;
            page-break-inside: avoid;
        }}
        
        .toc-link {{
            display: flex;
            justify-content: space-between;
            align-items: baseline;
            text-decoration: none;
            color: {options.color_text};
            padding: 0.3em 0;
            border-bottom: 1pt dotted #ccc;
        }}
        
        .toc-link:hover {{
            color: {options.color_accent};
        }}
        
        .toc-text {{
            flex: 1;
        }}
        
        .toc-page-num {{
            font-weight: 600;
            color: {options.color_accent};
            margin-left: 1em;
        }}
        
        .toc-level-2 {{
            margin-left: 1.5em;
            font-size: 0.95em;
        }}
        
        .toc-level-3 {{
            margin-left: 3em;
            font-size: 0.9em;
        }}
        
        /* === CONTENIDO EDITORIAL PRINCIPAL === */
        .main-content {{
            page-break-before: always;
        }}
        
        .editorial-title {{
            font-family: "{options.font_family_headings}", serif;
            font-size: {options.font_size_title}pt;
            font-weight: 900;
            color: {options.color_headings};
            text-align: center;
            margin: {options.chapter_spacing_before}em 0 {options.chapter_spacing_after}em 0;
            page-break-after: avoid;
            line-height: 1.1;
            letter-spacing: -0.02em;
        }}
        
        .editorial-chapter {{
            font-family: "{options.font_family_headings}", serif;
            font-size: {options.font_size_chapter}pt;
            font-weight: {options.font_weight_headings};
            color: {options.color_headings};
            margin: {options.chapter_spacing_before}em 0 {options.chapter_spacing_after}em 0;
            page-break-before: {'always' if options.enable_chapter_breaks else 'auto'};
            page-break-after: avoid;
            line-height: 1.2;
        }}
        
        .editorial-chapter::before {{
            content: counter(chapter, upper-roman) ". ";
            color: {options.color_accent};
            font-size: 0.8em;
            font-weight: 400;
            display: block;
            margin-bottom: 0.3em;
            text-transform: uppercase;
            letter-spacing: 0.2em;
        }}
        
        .editorial-section {{
            font-family: "{options.font_family_headings}", serif;
            font-size: {options.font_size_section}pt;
            font-weight: 600;
            color: {options.color_text};
            margin: {options.section_spacing_before}em 0 {options.section_spacing_after}em 0;
            page-break-after: avoid;
        }}
        
        .editorial-paragraph {{
            margin: {options.paragraph_spacing}em 0;
            text-indent: 1.5em;
            text-align: {'justify' if options.enable_justification else 'left'};
            orphans: {'3' if options.enable_orphan_control else 'auto'};
            widows: {'3' if options.enable_widow_control else 'auto'};
        }}
        
        .editorial-paragraph:first-of-type,
        .editorial-paragraph.first-paragraph {{
            text-indent: 0;
        }}
        
        /* === DROP CAPS EDITORIAL === */
        .editorial-paragraph.drop-caps::first-letter,
        .editorial-chapter + .editorial-paragraph::first-letter {{
            float: left;
            font-family: "{options.font_family_headings}", serif;
            font-size: 4em;
            line-height: 0.8;
            margin: 0.1em 0.1em 0 0;
            color: {options.color_accent};
            font-weight: 700;
            text-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }}
        
        /* === ELEMENTOS ESPECIALES === */
        .expression {{
            background: linear-gradient(135deg, #f7fafc 0%, #edf2f7 100%);
            border-left: 4pt solid {options.color_accent};
            padding: 1em 1.5em;
            margin: 1.5em 0;
            border-radius: 0 8pt 8pt 0;
            page-break-inside: avoid;
            font-weight: 500;
            box-shadow: 0 2pt 4pt rgba(0,0,0,0.1);
        }}
        
        .translation {{
            background: #f0fff4;
            border-left: 4pt solid #48bb78;
            padding: 0.8em 1.2em;
            margin: 1em 0;
            font-style: italic;
            border-radius: 0 6pt 6pt 0;
            page-break-inside: avoid;
        }}
        
        /* === CONTADORES EDITORIALES === */
        body {{
            counter-reset: chapter section page;
        }}
        
        .editorial-chapter {{
            counter-increment: chapter;
            counter-reset: section;
        }}
        
        .editorial-section {{
            counter-increment: section;
        }}
        
        /* === OPTIMIZACIÓN DE IMPRESIÓN === */
        @media print {{
            body {{
                -webkit-print-color-adjust: exact;
                print-color-adjust: exact;
            }}
            
            .page-break {{
                break-after: page;
            }}
            
            .expression,
            .translation {{
                break-inside: avoid;
            }}
        }}
        '''
    
    def _setup_advanced_document_properties(self,
                                          doc: Document,
                                          book_data: Dict,
                                          options: UltimateProfessionalOptions) -> None:
        """Configura propiedades avanzadas del documento Word."""
        
        core_props = doc.core_properties
        
        core_props.title = book_data.get('title', options.title)
        core_props.author = book_data.get('author', options.author)
        core_props.subject = options.subject
        core_props.keywords = '; '.join(options.keywords) if options.keywords else ''
        core_props.category = "Editorial Professional Document"
        core_props.comments = f"Generated with Buko AI Ultimate Professional Generator - Session: {self.session_id}"
        core_props.created = datetime.now()
        core_props.modified = datetime.now()
        core_props.language = options.language
        core_props.version = "1.0 Editorial Supreme"
    
    def _create_professional_styles(self,
                                  doc: Document,
                                  options: UltimateProfessionalOptions) -> None:
        """Crea estilos profesionales avanzados para Word."""
        
        # Estilo Normal (párrafos editoriales)
        normal_style = doc.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = options.font_family_body
        normal_font.size = Pt(options.font_size_body)
        normal_font.color.rgb = RGBColor.from_string(options.color_text.replace('#', ''))
        
        normal_paragraph = normal_style.paragraph_format
        normal_paragraph.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        normal_paragraph.line_spacing = options.line_height
        normal_paragraph.space_before = Pt(options.paragraph_spacing * 6)
        normal_paragraph.space_after = Pt(options.paragraph_spacing * 6)
        normal_paragraph.first_line_indent = Inches(0.5)
        
        if options.enable_justification:
            normal_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Estilo Título Editorial
        title_style = doc.styles.add_style('Editorial Title', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = options.font_family_headings
        title_font.size = Pt(options.font_size_title)
        title_font.bold = True
        title_font.color.rgb = RGBColor.from_string(options.color_headings.replace('#', ''))
        
        title_paragraph = title_style.paragraph_format
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_paragraph.space_before = Pt(options.chapter_spacing_before * 12)
        title_paragraph.space_after = Pt(options.chapter_spacing_after * 12)
        title_paragraph.keep_with_next = True
        
        # Estilo Capítulo Editorial
        chapter_style = doc.styles.add_style('Editorial Chapter', WD_STYLE_TYPE.PARAGRAPH)
        chapter_font = chapter_style.font
        chapter_font.name = options.font_family_headings
        chapter_font.size = Pt(options.font_size_chapter)
        chapter_font.bold = True
        chapter_font.color.rgb = RGBColor.from_string(options.color_headings.replace('#', ''))
        
        chapter_paragraph = chapter_style.paragraph_format
        chapter_paragraph.space_before = Pt(options.chapter_spacing_before * 12)
        chapter_paragraph.space_after = Pt(options.chapter_spacing_after * 12)
        chapter_paragraph.keep_with_next = True
        if options.enable_chapter_breaks:
            chapter_paragraph.page_break_before = True
        
        # Estilo Sección Editorial
        section_style = doc.styles.add_style('Editorial Section', WD_STYLE_TYPE.PARAGRAPH)
        section_font = section_style.font
        section_font.name = options.font_family_headings
        section_font.size = Pt(options.font_size_section)
        section_font.bold = True
        section_font.color.rgb = RGBColor.from_string(options.color_text.replace('#', ''))
        
        section_paragraph = section_style.paragraph_format
        section_paragraph.space_before = Pt(options.section_spacing_before * 12)
        section_paragraph.space_after = Pt(options.section_spacing_after * 12)
        section_paragraph.keep_with_next = True
    
    def _setup_advanced_page_layout(self, doc: Document, options: UltimateProfessionalOptions) -> None:
        """Configura layout avanzado de página."""
        section = doc.sections[0]
        
        # Configurar tamaño de página
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
        
        # Configurar márgenes
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.0)
        
        # Headers y footers si están habilitados
        if options.enable_headers:
            header = section.header
            header_para = header.paragraphs[0]
            header_para.text = options.title[:50] + "..." if len(options.title) > 50 else options.title
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if options.enable_footers and options.enable_page_numbers:
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_para.add_run().text = "— "
            
            # Agregar número de página
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            
            instrText = OxmlElement('w:instrText')
            instrText.text = "PAGE"
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            
            run = footer_para.runs[-1]
            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)
            
            footer_para.add_run().text = " —"
    
    def _add_professional_cover_page(self, doc: Document, book_data: Dict, options: UltimateProfessionalOptions) -> None:
        """Agrega página de portada profesional."""
        # Título principal
        title_para = doc.add_paragraph()
        title_para.style = 'Editorial Title'
        title_run = title_para.add_run(book_data.get('title', 'Título del Libro'))
        title_run.font.size = Pt(28)
        title_run.bold = True
        
        # Espaciado
        for _ in range(8):
            doc.add_paragraph()
        
        # Autor
        if book_data.get('author'):
            author_para = doc.add_paragraph()
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_run = author_para.add_run(f"Por {book_data['author']}")
            author_run.font.size = Pt(18)
            author_run.font.name = options.font_family_headings
        
        # Más espaciado
        for _ in range(10):
            doc.add_paragraph()
        
        # Información de publicación
        pub_para = doc.add_paragraph()
        pub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pub_run = pub_para.add_run(f"{options.publisher}\n{options.edition}\n{options.copyright_year}")
        pub_run.font.size = Pt(14)
        pub_run.font.italic = True
        
        # Salto de página
        doc.add_page_break()
    
    def _add_comprehensive_copyright_page(self, doc: Document, book_data: Dict, options: UltimateProfessionalOptions) -> None:
        """Agrega página de copyright comprehensive."""
        # Copyright notice
        copyright_para = doc.add_paragraph()
        copyright_run = copyright_para.add_run(f"Copyright © {options.copyright_year} {book_data.get('author', options.author)}")
        copyright_run.font.size = Pt(12)
        copyright_run.bold = True
        
        doc.add_paragraph()
        
        # Rights reserved
        rights_para = doc.add_paragraph()
        rights_text = ("Todos los derechos reservados. Ninguna parte de esta publicación puede ser "
                      "reproducida, distribuida o transmitida en cualquier forma o por cualquier medio, "
                      "incluyendo fotocopias, grabación u otros métodos electrónicos o mecánicos, "
                      "sin el permiso previo por escrito del editor.")
        rights_para.add_run(rights_text).font.size = Pt(10)
        
        doc.add_paragraph()
        
        # Publication info
        pub_info = doc.add_paragraph()
        pub_text = f"{options.edition}\nGenerado con tecnología de Inteligencia Artificial\n{options.publisher}"
        if options.isbn:
            pub_text += f"\nISBN: {options.isbn}"
        pub_info.add_run(pub_text).font.size = Pt(10)
        
        doc.add_page_break()
    
    def _add_professional_table_of_contents(self, doc: Document, processed_content: Dict, options: UltimateProfessionalOptions) -> None:
        """Agrega tabla de contenidos profesional sin duplicados."""
        # Título TOC
        toc_title = doc.add_paragraph()
        toc_title.style = 'Editorial Chapter'
        toc_title.add_run("Tabla de Contenidos")
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Entradas TOC (solo capítulos reales, sin duplicados)
        chapters = processed_content.get('chapters', [])
        page_start = 5 + (2 if options.include_copyright_page else 1)  # Ajustar por páginas anteriores
        
        for i, chapter in enumerate(chapters):
            toc_entry = doc.add_paragraph()
            # Calcular página basada en posición y contenido estimado
            estimated_page = page_start + (i * 2)  # Estimación más realista
            toc_run = toc_entry.add_run(f"{chapter['title']} {'.' * (50 - len(chapter['title']))} {estimated_page}")
            toc_run.font.size = Pt(12)
            
            # Agregar sangría para secciones
            para_format = toc_entry.paragraph_format
            if chapter.get('level', 2) > 2:
                para_format.left_indent = Inches(0.5)
        
        # Agregar secciones principales al TOC
        for section in processed_content.get('sections', [])[:5]:  # Solo las primeras 5 secciones
            if section.get('level', 3) == 3:  # Solo h3
                toc_entry = doc.add_paragraph()
                estimated_page = page_start + len(chapters) * 2 + section['position']
                toc_run = toc_entry.add_run(f"  {section['title']} {'.' * (45 - len(section['title']))} {estimated_page}")
                toc_run.font.size = Pt(10)
                para_format = toc_entry.paragraph_format
                para_format.left_indent = Inches(0.25)
            
        doc.add_page_break()
    
    def _add_formatted_main_content(self, doc: Document, processed_content: Dict, options: UltimateProfessionalOptions) -> None:
        """Agrega contenido principal formateado con estructura completa."""
        soup = BeautifulSoup(processed_content['raw_content'], 'html.parser')
        
        # Procesar todos los elementos en orden
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'div', 'ul', 'ol', 'li', 'table', 'img']):
            text_content = element.get_text().strip()
            element_classes = element.get('class', [])
            
            if not text_content and element.name not in ['img', 'table']:
                continue
            
            # Títulos principales
            if (element.name == 'h1' and any(cls in element_classes for cls in ['book-title', 'ebook-book-title'])):
                para = doc.add_paragraph()
                para.style = 'Editorial Title'
                para.add_run(text_content)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            # Capítulos
            elif (element.name == 'h2' or any(cls in element_classes for cls in ['chapter', 'ebook-chapter-title'])):
                if options.enable_chapter_breaks:
                    doc.add_page_break()
                para = doc.add_paragraph()
                para.style = 'Editorial Chapter'
                para.add_run(text_content)
                
            # Secciones
            elif element.name in ['h3', 'h4', 'h5', 'h6'] or any(cls in element_classes for cls in ['section', 'ebook-section']):
                para = doc.add_paragraph()
                para.style = 'Editorial Section'
                para.add_run(text_content)
                
            # Párrafos normales
            elif element.name == 'p':
                para = doc.add_paragraph()
                para.style = doc.styles['Normal']
                run = para.add_run(text_content)
                
                # Espaciado entre párrafos
                para_format = para.paragraph_format
                para_format.space_after = Pt(6)
                
            # Expresiones especiales
            elif any(cls in element_classes for cls in ['expression', 'ebook-expression']):
                para = doc.add_paragraph()
                para.style = doc.styles['Normal']
                run = para.add_run(text_content)
                run.bold = True
                run.italic = True
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            # Traducciones
            elif any(cls in element_classes for cls in ['translation', 'ebook-translation']):
                para = doc.add_paragraph()
                para.style = doc.styles['Normal']
                run = para.add_run(text_content)
                run.italic = True
                run.font.color.rgb = RGBColor(102, 102, 102)
                
            # Elementos de lista
            elif element.name in ['ul', 'ol']:
                # Agregar espacio antes de la lista
                doc.add_paragraph()
                
            elif element.name == 'li':
                para = doc.add_paragraph()
                para.style = doc.styles['Normal']
                # Agregar bullet manual
                run = para.add_run(f"• {text_content}")
                para_format = para.paragraph_format
                para_format.left_indent = Inches(0.5)
                
            # Divisores o contenedores
            elif element.name == 'div' and text_content:
                para = doc.add_paragraph()
                para.style = doc.styles['Normal']
                para.add_run(text_content)
    
    def _add_professional_index(self, doc: Document, processed_content: Dict, options: UltimateProfessionalOptions) -> None:
        """Agrega índice profesional."""
        # Título del índice
        index_title = doc.add_paragraph()
        index_title.style = 'Editorial Chapter'
        index_title.add_run("Índice")
        index_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Entradas del índice basadas en expresiones
        for expr in processed_content.get('expressions', []):
            index_entry = doc.add_paragraph()
            index_text = expr['text'][:50] + "..." if len(expr['text']) > 50 else expr['text']
            index_run = index_entry.add_run(f"{index_text} ........ {expr['number']}")
            index_run.font.size = Pt(10)
    
    def _add_about_author_page(self, doc: Document, book_data: Dict, options: UltimateProfessionalOptions) -> None:
        """Agrega página acerca del autor."""
        doc.add_page_break()
        
        # Título
        about_title = doc.add_paragraph()
        about_title.style = 'Editorial Chapter'
        about_title.add_run("Acerca del Autor")
        about_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Información del autor
        author_info = doc.add_paragraph()
        author_text = (f"{book_data.get('author', options.author)} es un autor dedicado a la creación "
                      f"de contenido educativo de alta calidad. Con la ayuda de tecnología de "
                      f"inteligencia artificial avanzada, ha logrado producir obras que combinan "
                      f"rigor académico con accesibilidad.")
        author_info.add_run(author_text).font.size = Pt(12)
        
        doc.add_paragraph()
        
        # Contacto
        contact_para = doc.add_paragraph()
        contact_para.add_run("Para más información, visite: www.buko-ai.com").font.size = Pt(10)
    
    def _generate_professional_html_document(self, processed_content: Dict, book_data: Dict, options: UltimateProfessionalOptions) -> str:
        """Genera documento HTML profesional."""
        html_filename = f"{book_data.get('title', 'documento')}_profesional.html"
        html_path = os.path.join(self.temp_dir, html_filename)
        
        # Generar HTML usando el método existente
        html_content = self._create_pdf_optimized_html(processed_content, book_data, options)
        
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        return html_path
    
    def _calculate_complexity_score(self, text: str) -> float:
        """Calcula un puntaje de complejidad del texto."""
        words = text.split()
        if not words:
            return 0.0
        
        # Factores de complejidad
        avg_word_length = sum(len(word) for word in words) / len(words)
        sentence_count = len(re.split(r'[.!?]+', text))
        avg_sentence_length = len(words) / max(1, sentence_count)
        
        # Puntaje simple de complejidad (0-100)
        complexity = min(100, (avg_word_length * 10) + (avg_sentence_length * 2))
        return round(complexity, 1)
    
    def _calculate_readability_score(self, text: str) -> float:
        """Calcula un puntaje de legibilidad simple."""
        words = text.split()
        if not words:
            return 0.0
        
        sentences = len(re.split(r'[.!?]+', text))
        avg_sentence_length = len(words) / max(1, sentences)
        
        # Puntaje simple de legibilidad (invertido de complejidad)
        readability = max(0, 100 - (avg_sentence_length * 3))
        return round(readability, 1)
    
    def _generate_comprehensive_document_info(self,
                                            processed_content: Dict,
                                            book_data: Dict,
                                            options: UltimateProfessionalOptions) -> Dict[str, Any]:
        """Genera información comprehensiva del documento generado."""
        
        return {
            'title': book_data.get('title', 'Documento Profesional'),
            'author': book_data.get('author', options.author),
            'publisher': options.publisher,
            'edition': options.edition,
            'isbn': options.isbn,
            'language': options.language,
            'copyright_year': options.copyright_year,
            'generation_date': datetime.now().isoformat(),
            'session_id': self.session_id,
            'quality_level': 'Editorial Supreme',
            'format_version': '2.0 Ultimate Professional',
            'statistics': processed_content.get('statistics', {}),
            'structure': {
                'chapters': len(processed_content.get('chapters', [])),
                'sections': len(processed_content.get('sections', [])),
                'expressions': len(processed_content.get('expressions', [])),
                'translations': len(processed_content.get('translations', []))
            },
            'features_enabled': {
                'professional_typography': True,
                'editorial_layout': True,
                'advanced_page_setup': True,
                'drop_caps': options.enable_drop_caps,
                'chapter_breaks': options.enable_chapter_breaks,
                'page_numbers': options.enable_page_numbers,
                'headers_footers': options.enable_headers and options.enable_footers,
                'table_of_contents': options.include_table_of_contents,
                'index': options.include_index,
                'justification': options.enable_justification,
                'hyphenation': options.enable_hyphenation,
                'orphan_widow_control': options.enable_orphan_control and options.enable_widow_control
            },
            'export_formats': {
                'word': True,
                'pdf': True,
                'html': True
            },
            'ready_for_publication': True,
            'commercial_grade': True
        }
    
    def cleanup(self):
        """Limpia archivos temporales."""
        try:
            import shutil
            shutil.rmtree(self.temp_dir)
        except Exception as e:
            print(f"Error cleaning temporary files: {e}")


def create_ultimate_professional_document(content: str,
                                        book_data: Dict[str, Any],
                                        user_options: Dict[str, Any] = None) -> Dict[str, Any]:
    """
    Función principal para crear documentos de calidad editorial suprema.
    
    Args:
        content: Contenido HTML/Markdown del libro
        book_data: Información del libro (título, autor, etc.)
        user_options: Opciones de formateo del usuario
        
    Returns:
        Dict con documentos generados e información completa
    """
    
    # Crear opciones profesionales supremas
    options = UltimateProfessionalOptions()
    
    # Aplicar configuraciones del usuario
    if user_options:
        # Mapear opciones de usuario a opciones profesionales
        if 'font_family' in user_options:
            options.font_family_body = user_options['font_family']
        if 'font_size_body' in user_options:
            options.font_size_body = int(user_options['font_size_body'])
        if 'line_spacing' in user_options:
            options.line_height = float(user_options['line_spacing'])
        if 'justify_text' in user_options:
            options.enable_justification = bool(user_options['justify_text'])
        if 'drop_caps' in user_options:
            options.enable_drop_caps = bool(user_options['drop_caps'])
        if 'page_numbers' in user_options:
            options.enable_page_numbers = bool(user_options['page_numbers'])
        if 'include_toc' in user_options:
            options.include_table_of_contents = bool(user_options['include_toc'])
        if 'include_cover' in user_options:
            options.include_cover_page = bool(user_options['include_cover'])
        
        # Configurar metadatos
        options.title = book_data.get('title', '')
        options.author = book_data.get('author', '')
        options.subject = book_data.get('genre', 'Literatura')
        
        if 'isbn' in user_options and user_options['isbn']:
            options.isbn = user_options['isbn']
    
    # Generar documento supremo
    generator = UltimateProfessionalGenerator()
    
    try:
        result = generator.generate_ultimate_professional_document(content, book_data, options)
        return result
    except Exception as e:
        generator.cleanup()
        raise e


if __name__ == "__main__":
    # Test del generador supremo
    print("🚀 INICIANDO TEST DEL GENERADOR SUPREMO")
    print("=" * 70)
    
    test_content = '''
    <h1 class="ebook-book-title">El Arte de la Programación Avanzada</h1>
    <h2 class="ebook-chapter-title">Capítulo 1: Fundamentos Revolucionarios</h2>
    <p class="ebook-paragraph">Este es el primer párrafo de nuestro libro profesional que demuestra la calidad editorial suprema que podemos alcanzar con las herramientas más avanzadas disponibles.</p>
    <div class="ebook-expression">**1. La programación es un arte que requiere precisión y creatividad.**</div>
    <div class="ebook-translation">Programming is an art that requires precision and creativity.</div>
    <h3 class="ebook-section">Sección 1.1: Principios Fundamentales</h3>
    <p class="ebook-paragraph">Los principios fundamentales de la programación avanzada incluyen la elegancia del código, la eficiencia algorítmica y la mantenibilidad a largo plazo.</p>
    '''
    
    test_book_data = {
        'title': 'El Arte de la Programación Avanzada',
        'author': 'Dr. Buko AI',
        'genre': 'Tecnología',
        'language': 'es'
    }
    
    test_user_options = {
        'font_family': 'Crimson Text',
        'font_size_body': 12,
        'line_spacing': 1.6,
        'justify_text': True,
        'drop_caps': True,
        'page_numbers': True,
        'include_toc': True,
        'include_cover': True,
        'isbn': '978-0-123456-78-9'
    }
    
    try:
        result = create_ultimate_professional_document(test_content, test_book_data, test_user_options)
        
        print("✅ GENERACIÓN SUPREMA COMPLETADA")
        print(f"📄 Documento Word: {result['word_path']}")
        print(f"📚 Documento PDF: {result['pdf_path']}")
        print(f"🌐 Documento HTML: {result['html_path']}")
        print(f"🏆 Calidad: {result['quality_level']}")
        print(f"📊 Páginas estimadas: {result['document_info']['statistics']['estimated_pages']}")
        print(f"📝 Palabras: {result['document_info']['statistics']['total_words']:,}")
        print(f"🎯 Listo para publicación: {result['ready_for_publication']}")
        
    except Exception as e:
        print(f"❌ Error en test: {str(e)}")
        import traceback
        traceback.print_exc()