"""
Professional HTML to DOCX Conversion Service
Convierte HTML generado por Claude a documentos DOCX profesionales con formateo automático.
"""

import re
import os
from typing import Dict, Any, Optional, List
from dataclasses import dataclass
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import logging

logger = logging.getLogger(__name__)

@dataclass
class ConversionOptions:
    """Opciones de conversión HTML a DOCX."""
    
    # Configuración de plataforma
    platform: str = "universal"
    
    # Opciones de documento
    page_width: float = 6.0  # pulgadas
    page_height: float = 9.0  # pulgadas
    margin_top: float = 1.0
    margin_bottom: float = 1.0
    margin_left: float = 1.0
    margin_right: float = 1.0
    
    # Opciones de tipografía
    font_family: str = "Times New Roman"
    font_size_body: int = 11
    font_size_h1: int = 18
    font_size_h2: int = 16
    font_size_h3: int = 14
    font_size_h4: int = 12
    
    # Opciones de espaciado
    line_spacing: float = 1.15
    paragraph_spacing_before: float = 6.0
    paragraph_spacing_after: float = 6.0
    
    # Opciones profesionales
    include_page_numbers: bool = True
    include_headers: bool = True
    use_professional_styles: bool = True
    create_table_of_contents: bool = True
    
    # Colores profesionales
    text_color: str = "#000000"
    heading_color: str = "#2c3e50"
    accent_color: str = "#3498db"
    warning_color: str = "#e74c3c"
    tip_color: str = "#27ae60"


class PlatformStyleConfig:
    """Configuraciones de estilo específicas por plataforma."""
    
    PLATFORM_CONFIGS = {
        "amazon_kdp": {
            "page_size": (6.0, 9.0),
            "margins": {"top": 1.0, "bottom": 1.0, "left": 1.0, "right": 1.0},
            "fonts": {"body": "Times New Roman", "heading": "Arial"},
            "font_sizes": {"body": 11, "h1": 18, "h2": 16, "h3": 14},
            "line_spacing": 1.15,
            "professional_elements": True
        },
        "google_play_books": {
            "page_size": (6.0, 9.0),
            "margins": {"top": 0.75, "bottom": 0.75, "left": 0.75, "right": 0.75},
            "fonts": {"body": "Times New Roman", "heading": "Arial"},
            "font_sizes": {"body": 12, "h1": 18, "h2": 16, "h3": 14},
            "line_spacing": 1.2,
            "professional_elements": True
        },
        "apple_books": {
            "page_size": (6.0, 9.0),
            "margins": {"top": 1.0, "bottom": 1.0, "left": 1.0, "right": 1.0},
            "fonts": {"body": "Times New Roman", "heading": "Arial"},
            "font_sizes": {"body": 11, "h1": 18, "h2": 16, "h3": 14},
            "line_spacing": 1.15,
            "professional_elements": True
        },
        "universal": {
            "page_size": (6.0, 9.0),
            "margins": {"top": 1.0, "bottom": 1.0, "left": 1.0, "right": 1.0},
            "fonts": {"body": "Times New Roman", "heading": "Arial"},
            "font_sizes": {"body": 11, "h1": 18, "h2": 16, "h3": 14},
            "line_spacing": 1.15,
            "professional_elements": True
        }
    }
    
    @classmethod
    def get_config(cls, platform: str) -> Dict[str, Any]:
        """Obtiene la configuración para una plataforma específica."""
        return cls.PLATFORM_CONFIGS.get(platform, cls.PLATFORM_CONFIGS["universal"])


class HTMLToDOCXConverter:
    """Convertidor profesional de HTML a DOCX."""
    
    def __init__(self, options: ConversionOptions = None):
        self.options = options or ConversionOptions()
        self.document = None
        self.current_list_level = 0
        self.table_of_contents_items = []
    
    def convert_html_to_docx(self, html_content: str, book_title: str = "", 
                           author: str = "Buko AI Editorial") -> Document:
        """
        Convierte HTML a documento DOCX profesional.
        
        Args:
            html_content: Contenido HTML a convertir
            book_title: Título del libro
            author: Autor del libro
            
        Returns:
            Documento DOCX profesional
        """
        try:
            logger.info(f"Iniciando conversión HTML a DOCX. Longitud HTML: {len(html_content)}")
            
            # Crear nuevo documento
            self.document = Document()
            
            # Configurar documento
            self._setup_document_properties(book_title, author)
            self._setup_page_layout()
            self._create_professional_styles()
            
            # Agregar páginas profesionales
            if self.options.use_professional_styles:
                self._add_title_page(book_title, author)
                self._add_copyright_page(author)
                
                if self.options.create_table_of_contents:
                    # Placeholder para TOC - se actualizará después de procesar el contenido
                    self._add_toc_placeholder()
            
            # Procesar HTML
            soup = BeautifulSoup(html_content, 'html.parser')
            self._process_html_elements(soup)
            
            # Actualizar tabla de contenidos
            if self.options.create_table_of_contents and self.table_of_contents_items:
                self._update_table_of_contents()
            
            logger.info("Conversión HTML a DOCX completada exitosamente")
            return self.document
            
        except Exception as e:
            logger.error(f"Error en conversión HTML a DOCX: {str(e)}")
            raise
    
    def _setup_document_properties(self, title: str, author: str):
        """Configura las propiedades del documento."""
        core_props = self.document.core_properties
        core_props.title = title
        core_props.author = author
        core_props.subject = "Libro generado con Buko AI"
        core_props.keywords = "ebook, professional, buko ai"
    
    def _setup_page_layout(self):
        """Configura el layout de página."""
        section = self.document.sections[0]
        
        # Tamaño de página
        section.page_width = Inches(self.options.page_width)
        section.page_height = Inches(self.options.page_height)
        
        # Márgenes
        section.top_margin = Inches(self.options.margin_top)
        section.bottom_margin = Inches(self.options.margin_bottom)
        section.left_margin = Inches(self.options.margin_left)
        section.right_margin = Inches(self.options.margin_right)
    
    def _create_professional_styles(self):
        """Crea estilos profesionales para el documento."""
        styles = self.document.styles
        
        # Estilo para párrafos normales
        if 'Professional Body' not in [s.name for s in styles]:
            body_style = styles.add_style('Professional Body', WD_STYLE_TYPE.PARAGRAPH)
            body_font = body_style.font
            body_font.name = self.options.font_family
            body_font.size = Pt(self.options.font_size_body)
            body_font.color.rgb = RGBColor.from_string(self.options.text_color.lstrip('#'))
            
            body_paragraph = body_style.paragraph_format
            body_paragraph.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            body_paragraph.line_spacing = self.options.line_spacing
            body_paragraph.space_before = Pt(self.options.paragraph_spacing_before)
            body_paragraph.space_after = Pt(self.options.paragraph_spacing_after)
            body_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Estilos para headings
        heading_configs = [
            ('Professional Heading 1', self.options.font_size_h1, True, 24, 18),
            ('Professional Heading 2', self.options.font_size_h2, True, 18, 12),
            ('Professional Heading 3', self.options.font_size_h3, True, 12, 6),
            ('Professional Heading 4', self.options.font_size_h4, True, 6, 6)
        ]
        
        for style_name, font_size, bold, space_before, space_after in heading_configs:
            if style_name not in [s.name for s in styles]:
                heading_style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                heading_font = heading_style.font
                heading_font.name = self.options.font_family
                heading_font.size = Pt(font_size)
                heading_font.bold = bold
                heading_font.color.rgb = RGBColor.from_string(self.options.heading_color.lstrip('#'))
                
                heading_paragraph = heading_style.paragraph_format
                heading_paragraph.space_before = Pt(space_before)
                heading_paragraph.space_after = Pt(space_after)
                heading_paragraph.keep_with_next = True
        
        # Estilos especiales
        special_styles = [
            ('Professional Example', self.options.accent_color, Cm(0.5)),
            ('Professional Tip', self.options.tip_color, Cm(0.5)),
            ('Professional Warning', self.options.warning_color, Cm(0.5))
        ]
        
        for style_name, color, indent in special_styles:
            if style_name not in [s.name for s in styles]:
                special_style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                special_font = special_style.font
                special_font.name = self.options.font_family
                special_font.size = Pt(self.options.font_size_body)
                
                special_paragraph = special_style.paragraph_format
                special_paragraph.left_indent = indent
                special_paragraph.space_before = Pt(6)
                special_paragraph.space_after = Pt(6)
    
    def _add_title_page(self, title: str, author: str):
        """Agrega página de título profesional."""
        # Título principal
        title_p = self.document.add_paragraph()
        title_run = title_p.add_run(title.upper())
        title_run.font.name = self.options.font_family
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.space_before = Pt(144)  # 2 pulgadas desde arriba
        title_p.space_after = Pt(72)
        
        # Autor
        author_p = self.document.add_paragraph()
        author_run = author_p.add_run(author)
        author_run.font.name = self.options.font_family
        author_run.font.size = Pt(16)
        author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_p.space_after = Pt(36)
        
        # Editorial
        editorial_p = self.document.add_paragraph()
        editorial_run = editorial_p.add_run("BUKO AI EDITORIAL")
        editorial_run.font.name = self.options.font_family
        editorial_run.font.size = Pt(12)
        editorial_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Salto de página
        self.document.add_page_break()
    
    def _add_copyright_page(self, author: str):
        """Agrega página de derechos de autor."""
        copyright_text = f"""Copyright © 2025 {author}

Todos los derechos reservados. Ninguna parte de esta publicación puede ser reproducida, distribuida o transmitida en cualquier forma o por cualquier medio, incluyendo fotocopias, grabación u otros métodos electrónicos o mecánicos, sin el permiso previo por escrito del editor, excepto en el caso de citas breves incorporadas en reseñas críticas y ciertos otros usos no comerciales permitidos por la ley de derechos de autor.

Primera edición: 2025

Publicado por Buko AI Editorial
Inteligencia Artificial Educativa

Para información sobre permisos o ediciones especiales, contacte:
info@bukoai.com"""
        
        copyright_p = self.document.add_paragraph(copyright_text)
        copyright_p.style = 'Professional Body'
        copyright_p.space_before = Pt(72)
        
        self.document.add_page_break()
    
    def _add_toc_placeholder(self):
        """Agrega placeholder para tabla de contenidos."""
        toc_title = self.document.add_paragraph("TABLA DE CONTENIDOS")
        toc_title.style = 'Professional Heading 1'
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Placeholder que se actualizará después
        self.toc_placeholder = self.document.add_paragraph("[Tabla de contenidos se generará automáticamente]")
        self.document.add_page_break()
    
    def _process_html_elements(self, soup: BeautifulSoup):
        """Procesa elementos HTML y los convierte a DOCX."""
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'ul', 'ol', 'li', 'div', 'strong', 'em', 'blockquote']):
            self._convert_element(element)
    
    def _convert_element(self, element):
        """Convierte un elemento HTML específico a DOCX."""
        tag_name = element.name
        
        if tag_name in ['h1', 'h2', 'h3', 'h4']:
            self._convert_heading(element, tag_name)
        elif tag_name == 'p':
            self._convert_paragraph(element)
        elif tag_name in ['ul', 'ol']:
            self._convert_list(element)
        elif tag_name == 'div':
            self._convert_div(element)
        elif tag_name == 'blockquote':
            self._convert_blockquote(element)
    
    def _convert_heading(self, element, tag_name: str):
        """Convierte headings HTML a estilos DOCX."""
        text = element.get_text().strip()
        if not text:
            return
        
        style_map = {
            'h1': 'Professional Heading 1',
            'h2': 'Professional Heading 2', 
            'h3': 'Professional Heading 3',
            'h4': 'Professional Heading 4'
        }
        
        paragraph = self.document.add_paragraph(text)
        paragraph.style = style_map[tag_name]
        
        # Agregar a tabla de contenidos
        level = int(tag_name[1])
        self.table_of_contents_items.append({
            'text': text,
            'level': level
        })
    
    def _convert_paragraph(self, element):
        """Convierte párrafos HTML a DOCX."""
        text = element.get_text().strip()
        if not text:
            return
        
        paragraph = self.document.add_paragraph()
        paragraph.style = 'Professional Body'
        
        # Procesar texto con formato
        self._process_inline_formatting(element, paragraph)
    
    def _convert_list(self, element):
        """Convierte listas HTML a DOCX."""
        list_items = element.find_all('li', recursive=False)
        
        for item in list_items:
            text = item.get_text().strip()
            if text:
                paragraph = self.document.add_paragraph()
                paragraph.style = 'List Bullet' if element.name == 'ul' else 'List Number'
                self._process_inline_formatting(item, paragraph)
    
    def _convert_div(self, element):
        """Convierte divs especiales HTML a DOCX."""
        class_name = element.get('class', [''])[0]
        text = element.get_text().strip()
        
        if not text:
            return
        
        style_map = {
            'example': 'Professional Example',
            'tip': 'Professional Tip',
            'warning': 'Professional Warning',
            'exercise': 'Professional Example',
            'case-study': 'Professional Example'
        }
        
        style = style_map.get(class_name, 'Professional Body')
        paragraph = self.document.add_paragraph(text)
        
        try:
            paragraph.style = style
        except:
            paragraph.style = 'Professional Body'
    
    def _convert_blockquote(self, element):
        """Convierte blockquotes HTML a DOCX."""
        text = element.get_text().strip()
        if text:
            paragraph = self.document.add_paragraph(f'"{text}"')
            paragraph.style = 'Professional Body'
            paragraph.paragraph_format.left_indent = Cm(1.0)
            paragraph.paragraph_format.right_indent = Cm(1.0)
    
    def _process_inline_formatting(self, element, paragraph):
        """Procesa formateo inline (strong, em, etc.)."""
        for content in element.contents:
            if hasattr(content, 'name'):
                if content.name == 'strong':
                    run = paragraph.add_run(content.get_text())
                    run.bold = True
                elif content.name == 'em':
                    run = paragraph.add_run(content.get_text())
                    run.italic = True
                else:
                    run = paragraph.add_run(content.get_text())
            else:
                # Texto plano
                text = str(content).strip()
                if text:
                    paragraph.add_run(text)
    
    def _update_table_of_contents(self):
        """Actualiza la tabla de contenidos con los headings encontrados."""
        if hasattr(self, 'toc_placeholder') and self.table_of_contents_items:
            # Limpiar placeholder
            self.toc_placeholder.clear()
            
            # Agregar entradas de TOC
            for item in self.table_of_contents_items:
                toc_text = "    " * (item['level'] - 1) + item['text']
                toc_paragraph = self.document.add_paragraph(toc_text)
                toc_paragraph.style = 'Professional Body'
                
                # Indentación basada en nivel
                indent = Cm(0.5 * (item['level'] - 1))
                toc_paragraph.paragraph_format.left_indent = indent
    
    def save_to_file(self, output_path: str) -> str:
        """Guarda el documento a un archivo."""
        try:
            self.document.save(output_path)
            logger.info(f"Documento DOCX guardado en: {output_path}")
            return output_path
        except Exception as e:
            logger.error(f"Error guardando DOCX: {str(e)}")
            raise


def convert_html_book_to_docx(html_content: str, output_path: str, 
                            book_title: str = "", author: str = "Buko AI Editorial",
                            platform: str = "universal") -> str:
    """
    Función de conveniencia para convertir HTML de libro a DOCX profesional.
    
    Args:
        html_content: Contenido HTML del libro
        output_path: Ruta donde guardar el archivo DOCX
        book_title: Título del libro
        author: Autor del libro
        platform: Plataforma de destino
        
    Returns:
        Ruta del archivo DOCX generado
    """
    try:
        # Configurar opciones basadas en plataforma
        platform_config = PlatformStyleConfig.get_config(platform)
        
        options = ConversionOptions(
            platform=platform,
            page_width=platform_config["page_size"][0],
            page_height=platform_config["page_size"][1],
            margin_top=platform_config["margins"]["top"],
            margin_bottom=platform_config["margins"]["bottom"],
            margin_left=platform_config["margins"]["left"],
            margin_right=platform_config["margins"]["right"],
            font_family=platform_config["fonts"]["body"],
            font_size_body=platform_config["font_sizes"]["body"],
            font_size_h1=platform_config["font_sizes"]["h1"],
            font_size_h2=platform_config["font_sizes"]["h2"],
            font_size_h3=platform_config["font_sizes"]["h3"],
            line_spacing=platform_config["line_spacing"],
            use_professional_styles=platform_config["professional_elements"]
        )
        
        # Convertir
        converter = HTMLToDOCXConverter(options)
        document = converter.convert_html_to_docx(html_content, book_title, author)
        
        # Guardar
        return converter.save_to_file(output_path)
        
    except Exception as e:
        logger.error(f"Error en conversión HTML a DOCX: {str(e)}")
        raise