"""
Generador Gratuito de Documentos Profesionales
Alternativa 100% gratuita a Aspose.Words usando python-docx, ReportLab y WeasyPrint.

Características principales:
- python-docx para documentos Word (.docx)
- ReportLab para PDFs profesionales 
- WeasyPrint para layouts complejos HTML→PDF
- Misma interfaz que AsposeProfessionalGenerator
- Sin costos de licencia
"""

import os
import tempfile
import logging
from typing import Dict, List, Optional, Any, Union
from dataclasses import dataclass, field
from enum import Enum
from datetime import datetime

# Free alternatives imports
try:
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_SECTION
    from docx.oxml.shared import OxmlElement, qn
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.units import inch
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    import weasyprint
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False

from app.models.book_generation import BookGeneration
from app.services.intelligent_content_generator import (
    IntelligentContentGenerator, 
    ContentType, 
    ContentGenerationRequest,
    BookAnalysis
)

# Import configuration classes from original service
from app.services.aspose_professional_generator import (
    PageFormat,
    FontFamily,
    TypographySettings,
    BookStructureSettings,
    ExportSettings,
    PageDimensions,
    AsposeDocumentConfiguration
)

logger = logging.getLogger(__name__)


class DocumentFormat(Enum):
    """Formatos de salida disponibles"""
    DOCX = "docx"
    PDF_REPORTLAB = "pdf_reportlab"
    PDF_WEASYPRINT = "pdf_weasyprint"
    HTML = "html"


class FreeDocumentGenerator:
    """
    Generador gratuito de documentos profesionales.
    Reemplaza completamente AsposeProfessionalGenerator sin costos de licencia.
    """
    
    def __init__(self):
        # Initialize content generator with fallback mode
        try:
            self.content_generator = IntelligentContentGenerator()
            self.content_generator_available = True
            logger.info("IntelligentContentGenerator initialized successfully")
        except Exception as e:
            logger.error(f"Error initializing IntelligentContentGenerator: {str(e)}")
            self.content_generator = None
            self.content_generator_available = False
            logger.warning("Using fallback content generation mode")
        
        # Check available generators
        self.available_formats = []
        if PYTHON_DOCX_AVAILABLE:
            self.available_formats.append(DocumentFormat.DOCX)
        if REPORTLAB_AVAILABLE:
            self.available_formats.append(DocumentFormat.PDF_REPORTLAB)
        if WEASYPRINT_AVAILABLE:
            self.available_formats.append(DocumentFormat.PDF_WEASYPRINT)
        
        logger.info(f"FreeDocumentGenerator initialized with formats: {[f.value for f in self.available_formats]}")
    
    def generate_professional_document(
        self, 
        book: BookGeneration, 
        config: AsposeDocumentConfiguration,
        output_path: Optional[str] = None,
        format: DocumentFormat = DocumentFormat.DOCX
    ) -> str:
        """
        Genera documento profesional completo usando alternativas gratuitas.
        Compatible con la interfaz original de AsposeProfessionalGenerator.
        """
        try:
            if format not in self.available_formats:
                raise ValueError(f"Format {format.value} not available. Available: {[f.value for f in self.available_formats]}")
            
            logger.info(f"Generating professional document for book: {book.title}")
            
            if format == DocumentFormat.DOCX:
                return self._generate_docx_document(book, config, output_path)
            elif format == DocumentFormat.PDF_REPORTLAB:
                return self._generate_pdf_reportlab(book, config, output_path)
            elif format == DocumentFormat.PDF_WEASYPRINT:
                return self._generate_pdf_weasyprint(book, config, output_path)
            else:
                raise ValueError(f"Unsupported format: {format}")
                
        except Exception as e:
            logger.error(f"Error generating professional document: {str(e)}")
            raise
    
    def generate_preview_document(
        self,
        book: BookGeneration,
        config: AsposeDocumentConfiguration,
        max_pages: int = 5
    ) -> str:
        """
        Genera preview limitado del documento para configuración rápida.
        Compatible con la interfaz original.
        """
        try:
            logger.info(f"Generating preview document (max {max_pages} pages)")
            
            # Use DOCX for preview by default
            if DocumentFormat.DOCX in self.available_formats:
                return self._generate_docx_preview(book, config, max_pages)
            elif DocumentFormat.PDF_REPORTLAB in self.available_formats:
                return self._generate_pdf_preview_reportlab(book, config, max_pages)
            else:
                raise RuntimeError("No document generators available")
                
        except Exception as e:
            logger.error(f"Error generating preview document: {str(e)}")
            raise
    
    def _generate_docx_document(
        self, 
        book: BookGeneration, 
        config: AsposeDocumentConfiguration,
        output_path: Optional[str] = None
    ) -> str:
        """Genera documento completo usando python-docx"""
        if not PYTHON_DOCX_AVAILABLE:
            raise RuntimeError("python-docx not available")
        
        # Create new document
        document = DocxDocument()
        
        # Configure page settings
        self._setup_page_settings_docx(document, config)
        
        # Generate content using intelligent generator
        book_analysis = self._get_book_analysis(book)
        
        # Add document sections based on configuration
        if config.structure.include_cover_page:
            self._add_cover_page_docx(document, book, config)
        
        if config.structure.include_title_page:
            self._add_title_page_docx(document, book, config)
        
        if config.structure.include_copyright_page:
            self._add_copyright_page_docx(document, book, config)
        
        if config.structure.include_dedication:
            self._add_dedication_docx(document, book, book_analysis)
        
        if config.structure.include_table_of_contents:
            self._add_table_of_contents_docx(document, book, book_analysis)
        
        if config.structure.include_prologue:
            self._add_prologue_docx(document, book, book_analysis)
        
        # Add main content
        self._add_main_content_docx(document, book, config, book_analysis)
        
        if config.structure.include_epilogue:
            self._add_epilogue_docx(document, book, book_analysis)
        
        if config.structure.include_about_author:
            self._add_about_author_docx(document, book, book_analysis)
        
        # Save document
        output_file = output_path or self._generate_output_filename(book, "docx")
        document.save(output_file)
        
        logger.info(f"DOCX document generated successfully: {output_file}")
        return output_file
    
    def _generate_docx_preview(
        self,
        book: BookGeneration,
        config: AsposeDocumentConfiguration,
        max_pages: int = 5
    ) -> str:
        """Genera preview limitado usando python-docx"""
        if not PYTHON_DOCX_AVAILABLE:
            raise RuntimeError("python-docx not available")
        
        # Create new document
        document = DocxDocument()
        
        # Configure page settings
        self._setup_page_settings_docx(document, config)
        
        # Generate limited content for preview
        book_analysis = self._get_book_analysis(book)
        
        # Add all configured sections for proper preview formatting
        if config.structure.include_cover_page:
            self._add_cover_page_docx(document, book, config)
        
        if config.structure.include_title_page:
            self._add_title_page_docx(document, book, config)
        
        if config.structure.include_copyright_page:
            self._add_copyright_page_docx(document, book, config)
        
        if config.structure.include_dedication:
            self._add_dedication_docx(document, book, book_analysis)
        
        if config.structure.include_table_of_contents:
            self._add_table_of_contents_docx(document, book, book_analysis)
        
        if config.structure.include_prologue:
            self._add_prologue_docx(document, book, book_analysis)
        
        # Add limited main content (first chapter or section)
        self._add_preview_content_docx(document, book, config, book_analysis, max_pages)
        
        if config.structure.include_epilogue:
            self._add_epilogue_docx(document, book, book_analysis)
        
        if config.structure.include_about_author:
            self._add_about_author_docx(document, book, book_analysis)
        
        # Save preview
        preview_file = self._generate_preview_filename(book, "docx")
        document.save(preview_file)
        
        return preview_file
    
    def _setup_page_settings_docx(self, document, config: AsposeDocumentConfiguration):
        """Configura las dimensiones y márgenes de página"""
        section = document.sections[0]
        
        # Set page dimensions based on format
        page_dims = self._get_page_dimensions(config.page_format, config.page_dimensions)
        
        section.page_width = Inches(page_dims['width'])
        section.page_height = Inches(page_dims['height'])
        section.top_margin = Inches(page_dims['margin_top'])
        section.bottom_margin = Inches(page_dims['margin_bottom'])
        section.left_margin = Inches(page_dims['margin_left'])
        section.right_margin = Inches(page_dims['margin_right'])
    
    def _add_cover_page_docx(self, document, book: BookGeneration, config):
        """Agrega página de portada profesional"""
        # Title
        title_paragraph = document.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(book.title)
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_run.font.name = config.typography.heading_font_family.value
        
        # Subtitle if available
        if hasattr(book, 'subtitle') and book.subtitle:
            subtitle_paragraph = document.add_paragraph()
            subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle_paragraph.add_run(book.subtitle)
            subtitle_run.font.size = Pt(16)
            subtitle_run.font.name = config.typography.heading_font_family.value
        
        # Author
        document.add_paragraph()  # Space
        author_paragraph = document.add_paragraph()
        author_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author_paragraph.add_run(config.author_name or "Autor")
        author_run.font.size = Pt(14)
        author_run.font.name = config.typography.body_font_family.value
        
        # Page break
        document.add_page_break()
    
    def _add_title_page_docx(self, document, book: BookGeneration, config):
        """Agrega página de título oficial"""
        title_paragraph = document.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(book.title)
        title_run.font.size = Pt(20)
        title_run.font.bold = True
        title_run.font.name = config.typography.heading_font_family.value
        
        document.add_page_break()
    
    def _add_copyright_page_docx(self, document, book: BookGeneration, config):
        """Agrega página de copyright"""
        copyright_text = f"""Copyright © {datetime.now().year} {config.author_name or 'Autor'}

Todos los derechos reservados. Ninguna parte de esta publicación puede ser reproducida, distribuida o transmitida en cualquier forma o por cualquier medio, incluyendo fotocopiado, grabación u otros métodos electrónicos o mecánicos, sin el permiso previo por escrito del editor, excepto en el caso de citas breves incorporadas en críticas y ciertos otros usos no comerciales permitidos por la ley de derechos de autor.

Generado con Buko AI - Sistema Inteligente de Generación de Libros
"""
        
        copyright_paragraph = document.add_paragraph(copyright_text)
        copyright_paragraph.runs[0].font.size = Pt(10)
        copyright_paragraph.runs[0].font.name = config.typography.body_font_family.value
        
        document.add_page_break()
    
    def _add_dedication_docx(self, document, book: BookGeneration, book_analysis: BookAnalysis):
        """Agrega dedicatoria generada inteligentemente"""
        if self.content_generator_available:
            try:
                request = ContentGenerationRequest(
                    content_type=ContentType.DEDICATION,
                    book_analysis=book_analysis,
                    specific_requirements={}
                )
                dedication = self.content_generator.generate_content(request)
            except Exception as e:
                logger.warning(f"Error generating dedication: {e}")
                dedication = "A todos aquellos que buscan conocimiento y crecimiento personal."
        else:
            dedication = "A todos aquellos que buscan conocimiento y crecimiento personal."
        
        # Add dedication centered and styled
        dedication_paragraph = document.add_paragraph()
        dedication_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dedication_run = dedication_paragraph.add_run(dedication)
        dedication_run.font.italic = True
        dedication_run.font.size = Pt(12)
        
        document.add_page_break()
    
    def _add_table_of_contents_docx(self, document, book: BookGeneration, book_analysis: BookAnalysis):
        """Agrega tabla de contenidos usando títulos detectados"""
        # Extract chapters for TOC
        chapters = self._extract_chapters_from_content(book.content)
        
        # Add TOC title
        toc_title = document.add_heading("Tabla de Contenidos", level=1)
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add TOC entries
        for chapter in chapters:
            toc_entry = document.add_paragraph()
            toc_entry.style = 'List Number'
            run = toc_entry.add_run(chapter.get('title', f"Capítulo {chapter['number']}"))
            run.font.size = Pt(12)
        
        document.add_page_break()
    
    def _add_prologue_docx(self, document, book: BookGeneration, book_analysis: BookAnalysis):
        """Agrega prólogo generado inteligentemente"""
        if self.content_generator_available:
            try:
                request = ContentGenerationRequest(
                    content_type=ContentType.PROLOGUE,
                    book_analysis=book_analysis,
                    specific_requirements={}
                )
                prologue = self.content_generator.generate_content(request)
            except Exception as e:
                logger.warning(f"Error generating prologue: {e}")
                prologue = "Este libro representa un viaje de descubrimiento y aprendizaje que esperamos sea de gran valor para el lector."
        else:
            prologue = "Este libro representa un viaje de descubrimiento y aprendizaje que esperamos sea de gran valor para el lector."
        
        # Add prologue title
        prologue_title = document.add_heading("Prólogo", level=1)
        prologue_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add prologue content
        prologue_paragraph = document.add_paragraph(prologue)
        prologue_paragraph.runs[0].font.size = Pt(11)
        
        document.add_page_break()
    
    def _add_epilogue_docx(self, document, book: BookGeneration, book_analysis: BookAnalysis):
        """Agrega epílogo generado inteligentemente"""
        if self.content_generator_available:
            try:
                request = ContentGenerationRequest(
                    content_type=ContentType.EPILOGUE,
                    book_analysis=book_analysis,
                    specific_requirements={}
                )
                epilogue = self.content_generator.generate_content(request)
            except Exception as e:
                logger.warning(f"Error generating epilogue: {e}")
                epilogue = "Esperamos que este recorrido haya sido enriquecedor y que los conocimientos compartidos sean de utilidad en su desarrollo personal y profesional."
        else:
            epilogue = "Esperamos que este recorrido haya sido enriquecedor y que los conocimientos compartidos sean de utilidad en su desarrollo personal y profesional."
        
        # Add epilogue title
        epilogue_title = document.add_heading("Epílogo", level=1)
        epilogue_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add epilogue content
        epilogue_paragraph = document.add_paragraph(epilogue)
        epilogue_paragraph.runs[0].font.size = Pt(11)
        
        document.add_page_break()
    
    def _add_about_author_docx(self, document, book: BookGeneration, book_analysis: BookAnalysis):
        """Agrega sección 'Acerca del Autor'"""
        if self.content_generator_available:
            try:
                request = ContentGenerationRequest(
                    content_type=ContentType.ABOUT_AUTHOR,
                    book_analysis=book_analysis,
                    specific_requirements={},
                    author_info={}
                )
                about_author = self.content_generator.generate_content(request)
            except Exception as e:
                logger.warning(f"Error generating about author: {e}")
                about_author = "El autor es un profesional dedicado a compartir conocimientos y experiencias a través de la escritura, con el objetivo de contribuir al crecimiento y desarrollo de los lectores."
        else:
            about_author = "El autor es un profesional dedicado a compartir conocimientos y experiencias a través de la escritura, con el objetivo de contribuir al crecimiento y desarrollo de los lectores."
        
        # Add about author title
        about_title = document.add_heading("Acerca del Autor", level=1)
        about_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add about author content
        about_paragraph = document.add_paragraph(about_author)
        about_paragraph.runs[0].font.size = Pt(11)
    
    def _add_main_content_docx(self, document, book: BookGeneration, config, book_analysis: BookAnalysis):
        """Agrega el contenido principal del libro con formateo profesional tipo Kindle"""
        # Convert HTML to professional format preserving structure
        professional_content = self._convert_html_to_professional_format(book.content)
        
        # Apply professional Kindle-style formatting
        self._apply_professional_kindle_formatting(document, professional_content, config)
    
    def _apply_professional_kindle_formatting(self, document, content: str, config):
        """Aplica formateo profesional estilo Kindle al contenido"""
        import re
        
        # Split content by chapter breaks first
        parts = re.split(r'=== CHAPTER_BREAK ===', content)
        
        for i, part in enumerate(parts):
            if not part.strip():
                continue
            
            # Extract chapter title if present
            chapter_match = re.search(r'^(.*?)\n=== END_CHAPTER_TITLE ===', part, re.DOTALL)
            if chapter_match:
                chapter_title = chapter_match.group(1).strip()
                remaining_content = part[chapter_match.end():].strip()
            else:
                # No explicit chapter title, create one
                chapter_title = f"Capítulo {i}" if i > 0 else "Introducción"
                remaining_content = part.strip()
            
            # Add chapter title with professional Kindle styling
            if chapter_title:
                chapter_heading = document.add_heading(chapter_title, level=1)
                chapter_heading.runs[0].font.name = config.typography.heading_font_family.value
                chapter_heading.runs[0].font.size = Pt(config.typography.chapter_title_size)
                chapter_heading.runs[0].font.bold = True
                chapter_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add space after chapter title (Kindle style)
                document.add_paragraph()
            
            # Process content sections
            sections = re.split(r'--- SECTION_HEADER ---', remaining_content)
            
            for j, section in enumerate(sections):
                if not section.strip():
                    continue
                
                # Extract section title if present
                section_match = re.search(r'^(.*?)\n--- END_SECTION ---', section, re.DOTALL)
                if section_match:
                    section_title = section_match.group(1).strip()
                    section_content = section[section_match.end():].strip()
                    
                    # Add section heading (Kindle style)
                    section_heading = document.add_heading(section_title, level=2)
                    section_heading.runs[0].font.name = config.typography.heading_font_family.value
                    section_heading.runs[0].font.size = Pt(config.typography.heading_font_size)
                    section_heading.runs[0].font.bold = True
                else:
                    section_content = section.strip()
                
                # Process paragraphs with professional formatting
                self._add_professional_paragraphs(document, section_content, config)
            
            # Add page break between chapters (Kindle style)
            if i < len(parts) - 1:
                document.add_page_break()
    
    def _add_professional_paragraphs(self, document, content: str, config):
        """Agrega párrafos con formateo profesional tipo Kindle"""
        import re
        
        # Handle special formatting blocks
        content = self._process_special_blocks(document, content, config)
        
        # Split into paragraphs
        paragraphs = content.split('\n\n')
        
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue
            
            # Create paragraph with professional formatting
            paragraph = document.add_paragraph()
            
            # Apply Kindle-style spacing and indentation
            paragraph.paragraph_format.first_line_indent = Inches(config.typography.first_line_indent / 72)
            paragraph.paragraph_format.space_before = Pt(config.typography.paragraph_spacing_before)
            paragraph.paragraph_format.space_after = Pt(config.typography.paragraph_spacing_after)
            paragraph.paragraph_format.line_spacing = config.typography.line_spacing
            
            # Process text with formatting markers
            self._add_formatted_text_runs(paragraph, para_text, config)
            
            # Apply justification (Kindle standard)
            if config.typography.justify_text:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    def _process_special_blocks(self, document, content: str, config):
        """Procesa bloques especiales como listas y citas"""
        import re
        
        # Handle quotes (blockquotes)
        def replace_quote(match):
            quote_text = match.group(1).strip()
            quote_para = document.add_paragraph()
            quote_para.style = 'Quote'
            run = quote_para.add_run(quote_text)
            run.font.italic = True
            run.font.size = Pt(config.typography.body_font_size - 1)
            return ""  # Remove from main content flow
        
        content = re.sub(r'--- QUOTE_START ---\n(.*?)\n--- QUOTE_END ---', replace_quote, content, flags=re.DOTALL)
        
        # Handle lists
        content = re.sub(r'--- LIST_START ---\n(.*?)\n--- LIST_END ---', r'\1', content, flags=re.DOTALL)
        content = re.sub(r'--- NUMBERED_LIST_START ---\n(.*?)\n--- NUMBERED_LIST_END ---', r'\1', content, flags=re.DOTALL)
        
        return content
    
    def _add_formatted_text_runs(self, paragraph, text: str, config):
        """Agrega texto con formateo avanzado (negrita, cursiva, etc.)"""
        import re
        
        # Split text by formatting markers
        parts = re.split(r'(\*\*BOLD\*\*.*?\*\*END_BOLD\*\*|\*ITALIC\*.*?\*END_ITALIC\*)', text)
        
        for part in parts:
            if not part:
                continue
            
            # Check for bold text
            bold_match = re.match(r'\*\*BOLD\*\*(.*?)\*\*END_BOLD\*\*', part)
            if bold_match:
                run = paragraph.add_run(bold_match.group(1))
                run.font.bold = True
                run.font.name = config.typography.body_font_family.value
                run.font.size = Pt(config.typography.body_font_size)
                continue
            
            # Check for italic text
            italic_match = re.match(r'\*ITALIC\*(.*?)\*END_ITALIC\*', part)
            if italic_match:
                run = paragraph.add_run(italic_match.group(1))
                run.font.italic = True
                run.font.name = config.typography.body_font_family.value
                run.font.size = Pt(config.typography.body_font_size)
                continue
            
            # Regular text
            run = paragraph.add_run(part)
            run.font.name = config.typography.body_font_family.value
            run.font.size = Pt(config.typography.body_font_size)
    
    def _add_preview_content_docx(self, document, book: BookGeneration, config, book_analysis: BookAnalysis, max_pages: int):
        """Agrega contenido limitado para preview con formateo profesional tipo Kindle"""
        # Convert HTML to professional format preserving structure
        professional_content = self._convert_html_to_professional_format(book.content)
        
        # Apply professional Kindle-style formatting for preview (limited content)
        self._apply_professional_kindle_formatting_preview(document, professional_content, config, max_pages)
    
    def _apply_professional_kindle_formatting_preview(self, document, content: str, config, max_pages: int):
        """Aplica formateo profesional estilo Kindle al preview (contenido limitado)"""
        import re
        
        # Split content by chapter breaks first
        parts = re.split(r'=== CHAPTER_BREAK ===', content)
        
        # Only process first chapter/section for preview
        if len(parts) > 1:
            first_part = parts[1]  # Skip empty first part
        else:
            first_part = parts[0] if parts else content
        
        # Extract chapter title if present
        chapter_match = re.search(r'^(.*?)\n=== END_CHAPTER_TITLE ===', first_part, re.DOTALL)
        if chapter_match:
            chapter_title = chapter_match.group(1).strip()
            remaining_content = first_part[chapter_match.end():].strip()
        else:
            # No explicit chapter title, create one
            chapter_title = "Capítulo 1"
            remaining_content = first_part.strip()
        
        # Add chapter title with professional Kindle styling
        if chapter_title:
            chapter_heading = document.add_heading(chapter_title, level=1)
            chapter_heading.runs[0].font.name = config.typography.heading_font_family.value
            chapter_heading.runs[0].font.size = Pt(config.typography.chapter_title_size)
            chapter_heading.runs[0].font.bold = True
            chapter_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add space after chapter title (Kindle style)
            document.add_paragraph()
        
        # Limit content for preview (estimate ~500 words per page)
        words_limit = max_pages * 500
        content_words = remaining_content.split()
        if len(content_words) > words_limit:
            limited_content = ' '.join(content_words[:words_limit])
            # Ensure we end at a sentence boundary
            last_sentence = limited_content.rfind('.')
            if last_sentence > len(limited_content) * 0.8:  # If found near the end
                limited_content = limited_content[:last_sentence + 1]
        else:
            limited_content = remaining_content
        
        # Process content sections for preview
        sections = re.split(r'--- SECTION_HEADER ---', limited_content)
        
        for j, section in enumerate(sections):
            if not section.strip():
                continue
            
            # Extract section title if present
            section_match = re.search(r'^(.*?)\n--- END_SECTION ---', section, re.DOTALL)
            if section_match:
                section_title = section_match.group(1).strip()
                section_content = section[section_match.end():].strip()
                
                # Add section heading (Kindle style)
                section_heading = document.add_heading(section_title, level=2)
                section_heading.runs[0].font.name = config.typography.heading_font_family.value
                section_heading.runs[0].font.size = Pt(config.typography.heading_font_size)
                section_heading.runs[0].font.bold = True
            else:
                section_content = section.strip()
            
            # Process paragraphs with professional formatting (limited for preview)
            self._add_professional_paragraphs_preview(document, section_content, config)
        
        # Add "preview continues" note with professional styling
        document.add_paragraph()  # Space before note
        continue_paragraph = document.add_paragraph()
        continue_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        continue_run = continue_paragraph.add_run("... continúa en la versión completa ...")
        continue_run.font.italic = True
        continue_run.font.size = Pt(10)
        continue_run.font.color.rgb = None  # Use default color but make it subtle
    
    def _add_professional_paragraphs_preview(self, document, content: str, config):
        """Agrega párrafos para preview con formateo profesional tipo Kindle (limitado)"""
        import re
        
        # Handle special formatting blocks
        content = self._process_special_blocks(document, content, config)
        
        # Split into paragraphs and limit for preview
        paragraphs = content.split('\n\n')
        
        # Limit to reasonable number of paragraphs for preview
        max_paragraphs = 15
        
        for i, para_text in enumerate(paragraphs[:max_paragraphs]):
            para_text = para_text.strip()
            if not para_text:
                continue
            
            # Create paragraph with professional formatting
            paragraph = document.add_paragraph()
            
            # Apply Kindle-style spacing and indentation
            paragraph.paragraph_format.first_line_indent = Inches(config.typography.first_line_indent / 72)
            paragraph.paragraph_format.space_before = Pt(config.typography.paragraph_spacing_before)
            paragraph.paragraph_format.space_after = Pt(config.typography.paragraph_spacing_after)
            paragraph.paragraph_format.line_spacing = config.typography.line_spacing
            
            # Process text with formatting markers
            self._add_formatted_text_runs(paragraph, para_text, config)
            
            # Apply justification (Kindle standard)
            if config.typography.justify_text:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    def _convert_html_to_professional_format(self, content: str) -> str:
        """Convierte HTML a formato profesional preservando estructura para libro tipo Kindle"""
        import re
        
        if not content:
            return ""
        
        # Step 1: Preserve and enhance structure for professional formatting
        professional_content = content
        
        # Detect and format chapter headers (H1, H2) as proper chapter breaks
        professional_content = re.sub(
            r'<h1[^>]*>(.*?)</h1>', 
            r'\n\n=== CHAPTER_BREAK ===\n\1\n=== END_CHAPTER_TITLE ===\n\n', 
            professional_content, 
            flags=re.DOTALL | re.IGNORECASE
        )
        
        # Subheadings (H2-H6) become section headers
        professional_content = re.sub(
            r'<h([2-6])[^>]*>(.*?)</h\1>', 
            r'\n\n--- SECTION_HEADER ---\n\2\n--- END_SECTION ---\n\n', 
            professional_content, 
            flags=re.DOTALL | re.IGNORECASE
        )
        
        # Enhanced paragraph handling for Kindle-style formatting
        professional_content = re.sub(
            r'<p[^>]*>(.*?)</p>', 
            r'\n\n\1\n\n', 
            professional_content, 
            flags=re.DOTALL | re.IGNORECASE
        )
        
        # Preserve important formatting markers for later processing
        professional_content = re.sub(r'<(strong|b)[^>]*>(.*?)</\1>', r'**BOLD**\2**END_BOLD**', professional_content, flags=re.DOTALL | re.IGNORECASE)
        professional_content = re.sub(r'<(em|i)[^>]*>(.*?)</\1>', r'*ITALIC*\2*END_ITALIC*', professional_content, flags=re.DOTALL | re.IGNORECASE)
        
        # Enhanced list handling for professional appearance
        professional_content = re.sub(r'<ul[^>]*>', r'\n--- LIST_START ---\n', professional_content, flags=re.IGNORECASE)
        professional_content = re.sub(r'</ul>', r'\n--- LIST_END ---\n\n', professional_content, flags=re.IGNORECASE)
        professional_content = re.sub(r'<ol[^>]*>', r'\n--- NUMBERED_LIST_START ---\n', professional_content, flags=re.IGNORECASE)
        professional_content = re.sub(r'</ol>', r'\n--- NUMBERED_LIST_END ---\n\n', professional_content, flags=re.IGNORECASE)
        professional_content = re.sub(r'<li[^>]*>(.*?)</li>', r'• \1\n', professional_content, flags=re.DOTALL | re.IGNORECASE)
        
        # Blockquotes for professional styling
        professional_content = re.sub(r'<blockquote[^>]*>(.*?)</blockquote>', r'\n--- QUOTE_START ---\n\1\n--- QUOTE_END ---\n\n', professional_content, flags=re.DOTALL | re.IGNORECASE)
        
        # Clean remaining HTML tags
        professional_content = re.sub(r'<[^>]+>', '', professional_content)
        
        # Decode HTML entities
        html_entities = {
            '&amp;': '&', '&lt;': '<', '&gt;': '>', '&quot;': '"', '&#39;': "'", '&nbsp;': ' ',
            '&aacute;': 'á', '&eacute;': 'é', '&iacute;': 'í', '&oacute;': 'ó', '&uacute;': 'ú', '&ntilde;': 'ñ',
            '&Aacute;': 'Á', '&Eacute;': 'É', '&Iacute;': 'Í', '&Oacute;': 'Ó', '&Uacute;': 'Ú', '&Ntilde;': 'Ñ',
            '&auml;': 'ä', '&ouml;': 'ö', '&uuml;': 'ü', '&Auml;': 'Ä', '&Ouml;': 'Ö', '&Uuml;': 'Ü', '&szlig;': 'ß'
        }
        
        for entity, char in html_entities.items():
            professional_content = professional_content.replace(entity, char)
        
        # Professional text cleanup for Kindle-quality formatting
        professional_content = re.sub(r'\n\s*\n\s*\n+', '\n\n', professional_content)  # Max 2 consecutive newlines
        professional_content = re.sub(r'[ \t]+', ' ', professional_content)  # Multiple spaces to single space
        professional_content = re.sub(r' +\n', '\n', professional_content)  # Remove trailing spaces
        professional_content = professional_content.strip()
        
        return professional_content
    
    def _clean_html_content(self, content: str) -> str:
        """Legacy method - redirects to professional formatter"""
        return self._convert_html_to_professional_format(content)
    
    def _extract_chapters_from_content(self, content: str) -> List[Dict[str, str]]:
        """Extrae capítulos del contenido del libro"""
        import re
        
        # Try to detect chapter patterns
        chapter_patterns = [
            r'(?i)^capítulo\s+(\d+)[\:\.]?\s*(.*)$',
            r'(?i)^chapter\s+(\d+)[\:\.]?\s*(.*)$',
            r'(?i)^(\d+)[\:\.]?\s+(.*)$'
        ]
        
        lines = content.split('\n')
        chapters = []
        current_chapter = None
        current_content = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if this line is a chapter header
            is_chapter_header = False
            for pattern in chapter_patterns:
                match = re.match(pattern, line)
                if match:
                    # Save previous chapter
                    if current_chapter is not None:
                        current_chapter['content'] = '\n\n'.join(current_content)
                        chapters.append(current_chapter)
                    
                    # Start new chapter
                    chapter_num = len(chapters) + 1
                    current_chapter = {
                        'number': chapter_num,
                        'title': match.group(2) if len(match.groups()) > 1 else f"Capítulo {chapter_num}"
                    }
                    current_content = []
                    is_chapter_header = True
                    break
            
            if not is_chapter_header:
                current_content.append(line)
        
        # Add last chapter
        if current_chapter is not None:
            current_chapter['content'] = '\n\n'.join(current_content)
            chapters.append(current_chapter)
        
        # If no chapters detected, create single chapter
        if not chapters:
            chapters = [{
                'number': 1,
                'title': 'Contenido Principal',
                'content': content
            }]
        
        return chapters
    
    def _get_book_analysis(self, book: BookGeneration) -> BookAnalysis:
        """Obtiene análisis del libro usando IntelligentContentGenerator"""
        if self.content_generator_available:
            try:
                return self.content_generator.analyze_book_content(book)
            except Exception as e:
                logger.error(f"Error analyzing book content: {e}")
                return self._get_fallback_analysis(book)
        else:
            return self._get_fallback_analysis(book)
    
    def _get_fallback_analysis(self, book: BookGeneration) -> BookAnalysis:
        """Análisis de respaldo básico"""
        from app.services.intelligent_content_generator import BookGenre, BookAnalysis
        
        return BookAnalysis(
            genre=BookGenre.NON_FICTION,
            main_themes=["conocimiento", "aprendizaje", "desarrollo"],
            tone="profesional",
            target_audience="adultos",
            language_style="informativo",
            key_concepts=["conceptos clave", "ideas principales"],
            chapter_structure=[],
            estimated_reading_level="intermedio",
            cultural_context="general"
        )
    
    def _get_page_dimensions(self, page_format: PageFormat, custom_dimensions: Optional[PageDimensions] = None) -> Dict[str, float]:
        """Obtiene dimensiones de página en pulgadas"""
        if page_format == PageFormat.CUSTOM and custom_dimensions:
            return {
                'width': custom_dimensions.width / 72,  # Convert points to inches
                'height': custom_dimensions.height / 72,
                'margin_top': custom_dimensions.margin_top / 72,
                'margin_bottom': custom_dimensions.margin_bottom / 72,
                'margin_left': custom_dimensions.margin_left / 72,
                'margin_right': custom_dimensions.margin_right / 72
            }
        
        # Standard page formats in inches
        formats = {
            PageFormat.POCKET: {'width': 4.25, 'height': 6.87},
            PageFormat.MASS_MARKET: {'width': 4.25, 'height': 7.0},
            PageFormat.TRADE_PAPERBACK: {'width': 6.0, 'height': 9.0},
            PageFormat.HARDCOVER: {'width': 6.14, 'height': 9.21},
            PageFormat.LARGE_FORMAT: {'width': 8.5, 'height': 11.0},
            PageFormat.SQUARE: {'width': 8.0, 'height': 8.0}
        }
        
        dimensions = formats.get(page_format, formats[PageFormat.TRADE_PAPERBACK])
        dimensions.update({
            'margin_top': 1.0,
            'margin_bottom': 1.0,
            'margin_left': 1.0,
            'margin_right': 1.0
        })
        
        return dimensions
    
    def _generate_output_filename(self, book: BookGeneration, extension: str) -> str:
        """Genera nombre de archivo de salida"""
        safe_title = "".join(c for c in book.title if c.isalnum() or c in (' ', '-', '_')).strip()
        safe_title = safe_title.replace(' ', '_')
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"/tmp/{safe_title}_professional_{timestamp}.{extension}"
    
    def _generate_preview_filename(self, book: BookGeneration, extension: str) -> str:
        """Genera nombre de archivo para preview"""
        safe_title = "".join(c for c in book.title if c.isalnum() or c in (' ', '-', '_')).strip()
        safe_title = safe_title.replace(' ', '_')
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"/tmp/preview_{safe_title}_{timestamp}.{extension}"
    
    # Additional methods for other formats (PDF, etc.) can be added here
    def _generate_pdf_reportlab(self, book: BookGeneration, config: AsposeDocumentConfiguration, output_path: Optional[str] = None) -> str:
        """Genera PDF usando ReportLab (implementación futura)"""
        raise NotImplementedError("PDF generation with ReportLab coming soon")
    
    def _generate_pdf_weasyprint(self, book: BookGeneration, config: AsposeDocumentConfiguration, output_path: Optional[str] = None) -> str:
        """Genera PDF usando WeasyPrint (implementación futura)"""
        raise NotImplementedError("PDF generation with WeasyPrint coming soon")


# Convenience function to maintain compatibility
def get_free_document_generator() -> FreeDocumentGenerator:
    """Función de conveniencia para obtener una instancia del generador gratuito"""
    return FreeDocumentGenerator()