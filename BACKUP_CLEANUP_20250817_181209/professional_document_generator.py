"""
Professional Document Generator - Salto Monumental
Genera documentos Word y PDF con formateo profesional REAL usando python-docx
"""

import os
import re
import tempfile
from typing import Dict, List, Any, Optional, Tuple
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.section import WD_ORIENT, WD_SECTION
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("⚠️  python-docx no está instalado. Instalar con: pip install python-docx")
    Document = None

try:
    import subprocess
    LIBREOFFICE_AVAILABLE = True
except ImportError:
    LIBREOFFICE_AVAILABLE = False

from bs4 import BeautifulSoup


@dataclass
class ProfessionalDocumentOptions:
    """Opciones profesionales para generación de documentos."""
    
    # TIPOGRAFÍA PROFESIONAL
    font_family: str = "Crimson Text"
    font_size_body: int = 12
    font_size_title: int = 18
    font_size_chapter: int = 16
    font_size_section: int = 14
    
    # ESPACIADO PROFESIONAL
    line_spacing: float = 1.15
    paragraph_spacing_before: int = 6
    paragraph_spacing_after: int = 6
    chapter_spacing_before: int = 24
    chapter_spacing_after: int = 12
    
    # MÁRGENES PROFESIONALES (en pulgadas)
    margin_top: float = 1.0
    margin_bottom: float = 1.0
    margin_left: float = 1.25
    margin_right: float = 1.0
    
    # CONFIGURACIÓN DE PÁGINA
    page_width: float = 8.5  # Letter size
    page_height: float = 11.0
    orientation: str = "portrait"  # portrait, landscape
    
    # ELEMENTOS PROFESIONALES
    include_header: bool = True
    include_footer: bool = True
    include_page_numbers: bool = True
    include_toc: bool = True
    include_cover_page: bool = True
    
    # ESTILOS AVANZADOS
    justify_text: bool = True
    hyphenation: bool = True
    first_line_indent: float = 0.5
    
    # COLORES PROFESIONALES
    text_color: str = "#000000"
    chapter_color: str = "#1f2937"
    accent_color: str = "#3b82f6"
    
    # METADATOS
    author: str = ""
    title: str = ""
    subject: str = ""
    keywords: str = ""


class ProfessionalDocumentGenerator:
    """Generador de documentos profesionales con formateo REAL."""
    
    def __init__(self):
        if Document is None:
            raise ImportError("python-docx es requerido para generar documentos profesionales")
        
        self.temp_dir = tempfile.mkdtemp()
        
    def generate_professional_document(self, 
                                     content: str, 
                                     book_data: Dict[str, Any],
                                     options: ProfessionalDocumentOptions) -> Dict[str, str]:
        """
        Genera documento profesional Word y PDF.
        
        Returns:
            Dict con paths al documento Word y PDF generados
        """
        
        # Crear documento Word
        doc = Document()
        
        # Configurar documento
        self._setup_document_properties(doc, book_data, options)
        self._setup_document_styles(doc, options)
        self._setup_page_layout(doc, options)
        
        # Generar contenido
        self._add_cover_page(doc, book_data, options)
        self._add_table_of_contents(doc, options)
        self._add_formatted_content(doc, content, options)
        
        # Guardar documento Word
        word_path = os.path.join(self.temp_dir, f"{book_data.get('title', 'book')}.docx")
        doc.save(word_path)
        
        # Convertir a PDF
        pdf_path = self._convert_to_pdf(word_path)
        
        return {
            'word_path': word_path,
            'pdf_path': pdf_path,
            'temp_dir': self.temp_dir
        }
    
    def _setup_document_properties(self, doc: Document, book_data: Dict, options: ProfessionalDocumentOptions):
        """Configura propiedades del documento."""
        core_props = doc.core_properties
        
        core_props.title = book_data.get('title', options.title)
        core_props.author = book_data.get('author', options.author)
        core_props.subject = options.subject
        core_props.keywords = options.keywords
        core_props.created = datetime.now()
        core_props.modified = datetime.now()
        core_props.category = "Libro Profesional Buko AI"
        core_props.comments = "Generado con tecnología de IA profesional"
    
    def _setup_document_styles(self, doc: Document, options: ProfessionalDocumentOptions):
        """Configura estilos profesionales del documento."""
        
        # Estilo Normal (párrafos)
        normal_style = doc.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = options.font_family
        normal_font.size = Pt(options.font_size_body)
        normal_font.color.rgb = RGBColor.from_string(options.text_color.replace('#', ''))
        
        normal_paragraph = normal_style.paragraph_format
        normal_paragraph.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        normal_paragraph.line_spacing = options.line_spacing
        normal_paragraph.space_before = Pt(options.paragraph_spacing_before)
        normal_paragraph.space_after = Pt(options.paragraph_spacing_after)
        normal_paragraph.first_line_indent = Inches(options.first_line_indent)
        
        if options.justify_text:
            normal_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Estilo Título Principal
        if 'Title' not in [style.name for style in doc.styles]:
            title_style = doc.styles.add_style('Professional Title', WD_STYLE_TYPE.PARAGRAPH)
        else:
            title_style = doc.styles['Title']
            
        title_font = title_style.font
        title_font.name = options.font_family
        title_font.size = Pt(options.font_size_title)
        title_font.bold = True
        title_font.color.rgb = RGBColor.from_string(options.chapter_color.replace('#', ''))
        
        title_paragraph = title_style.paragraph_format
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_paragraph.space_before = Pt(24)
        title_paragraph.space_after = Pt(24)
        
        # Estilo Capítulo
        chapter_style = doc.styles.add_style('Professional Chapter', WD_STYLE_TYPE.PARAGRAPH)
        chapter_font = chapter_style.font
        chapter_font.name = options.font_family
        chapter_font.size = Pt(options.font_size_chapter)
        chapter_font.bold = True
        chapter_font.color.rgb = RGBColor.from_string(options.chapter_color.replace('#', ''))
        
        chapter_paragraph = chapter_style.paragraph_format
        chapter_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        chapter_paragraph.space_before = Pt(options.chapter_spacing_before)
        chapter_paragraph.space_after = Pt(options.chapter_spacing_after)
        chapter_paragraph.keep_with_next = True
        
        # Estilo Sección
        section_style = doc.styles.add_style('Professional Section', WD_STYLE_TYPE.PARAGRAPH)
        section_font = section_style.font
        section_font.name = options.font_family
        section_font.size = Pt(options.font_size_section)
        section_font.bold = True
        
        section_paragraph = section_style.paragraph_format
        section_paragraph.space_before = Pt(12)
        section_paragraph.space_after = Pt(6)
        section_paragraph.keep_with_next = True
        
        # Estilo Expresión (destacado)
        expression_style = doc.styles.add_style('Professional Expression', WD_STYLE_TYPE.PARAGRAPH)
        expression_font = expression_style.font
        expression_font.name = options.font_family
        expression_font.size = Pt(options.font_size_body)
        expression_font.bold = True
        expression_font.color.rgb = RGBColor.from_string(options.accent_color.replace('#', ''))
        
        expression_paragraph = expression_style.paragraph_format
        expression_paragraph.left_indent = Inches(0.5)
        expression_paragraph.space_before = Pt(6)
        expression_paragraph.space_after = Pt(6)
        
        # Estilo Traducción
        translation_style = doc.styles.add_style('Professional Translation', WD_STYLE_TYPE.PARAGRAPH)
        translation_font = translation_style.font
        translation_font.name = options.font_family
        translation_font.size = Pt(options.font_size_body - 1)
        translation_font.italic = True
        
        translation_paragraph = translation_style.paragraph_format
        translation_paragraph.left_indent = Inches(0.75)
        translation_paragraph.space_before = Pt(3)
        translation_paragraph.space_after = Pt(3)
    
    def _setup_page_layout(self, doc: Document, options: ProfessionalDocumentOptions):
        """Configura el layout de página profesional."""
        
        section = doc.sections[0]
        
        # Configurar tamaño de página
        section.page_width = Inches(options.page_width)
        section.page_height = Inches(options.page_height)
        
        # Configurar márgenes
        section.top_margin = Inches(options.margin_top)
        section.bottom_margin = Inches(options.margin_bottom)
        section.left_margin = Inches(options.margin_left)
        section.right_margin = Inches(options.margin_right)
        
        # Configurar header y footer
        if options.include_header:
            header = section.header
            header_para = header.paragraphs[0]
            header_para.text = f"{options.title}"
            header_para.style = doc.styles['Normal']
            header_font = header_para.runs[0].font
            header_font.size = Pt(10)
            header_font.italic = True
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if options.include_footer and options.include_page_numbers:
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Agregar número de página
            run = footer_para.runs[0] if footer_para.runs else footer_para.add_run()
            run.font.size = Pt(10)
            
            # XML para número de página
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            
            instrText = OxmlElement('w:instrText')
            instrText.text = "PAGE"
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            
            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)
    
    def _add_cover_page(self, doc: Document, book_data: Dict, options: ProfessionalDocumentOptions):
        """Agrega página de portada profesional."""
        if not options.include_cover_page:
            return
        
        # Título principal
        title_para = doc.add_paragraph()
        title_para.style = 'Professional Title'
        title_run = title_para.add_run(book_data.get('title', 'Título del Libro'))
        title_run.font.size = Pt(24)
        title_run.bold = True
        
        # Espaciado
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Autor
        if book_data.get('author'):
            author_para = doc.add_paragraph()
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_run = author_para.add_run(f"Por {book_data['author']}")
            author_run.font.size = Pt(16)
            author_run.font.name = options.font_family
        
        # Espaciado adicional
        for _ in range(10):
            doc.add_paragraph()
        
        # Información de publicación
        pub_para = doc.add_paragraph()
        pub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pub_run = pub_para.add_run("Generado con Buko AI\nTecnología de Inteligencia Artificial")
        pub_run.font.size = Pt(12)
        pub_run.font.italic = True
        
        # Fecha
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(datetime.now().strftime("%B %Y"))
        date_run.font.size = Pt(12)
        
        # Salto de página
        doc.add_page_break()
    
    def _add_table_of_contents(self, doc: Document, options: ProfessionalDocumentOptions):
        """Agrega tabla de contenidos profesional."""
        if not options.include_toc:
            return
        
        # Título TOC
        toc_title = doc.add_paragraph()
        toc_title.style = 'Professional Chapter'
        toc_title.add_run("Tabla de Contenidos")
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Nota: En Word real, la TOC se genera automáticamente
        # Aquí agregamos un placeholder que se puede actualizar en Word
        toc_note = doc.add_paragraph()
        toc_note.add_run("(La tabla de contenidos se generará automáticamente al abrir en Microsoft Word)")
        toc_note.style = doc.styles['Normal']
        toc_note.runs[0].font.italic = True
        toc_note.runs[0].font.size = Pt(10)
        
        doc.add_page_break()
    
    def _add_formatted_content(self, doc: Document, content: str, options: ProfessionalDocumentOptions):
        """Agrega contenido formateado profesionalmente."""
        
        # Parsear HTML
        soup = BeautifulSoup(content, 'html.parser')
        
        # Procesar elementos
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'div']):
            self._process_element(doc, element, options)
    
    def _process_element(self, doc: Document, element, options: ProfessionalDocumentOptions):
        """Procesa un elemento HTML y lo convierte a formato Word profesional."""
        
        text_content = element.get_text().strip()
        if not text_content:
            return
        
        # Detectar tipo de elemento
        if element.name == 'h1' or 'book-title' in element.get('class', []):
            # Título principal
            para = doc.add_paragraph()
            para.style = 'Professional Title'
            para.add_run(text_content)
            
        elif element.name == 'h2' or 'chapter' in element.get('class', []):
            # Capítulo
            para = doc.add_paragraph()
            para.style = 'Professional Chapter'
            para.add_run(text_content)
            
        elif element.name == 'h3' or 'section' in element.get('class', []):
            # Sección
            para = doc.add_paragraph()
            para.style = 'Professional Section'
            para.add_run(text_content)
            
        elif 'expression' in element.get('class', []):
            # Expresión numerada
            para = doc.add_paragraph()
            para.style = 'Professional Expression'
            para.add_run(text_content)
            
        elif 'translation' in element.get('class', []):
            # Traducción
            para = doc.add_paragraph()
            para.style = 'Professional Translation'
            para.add_run(text_content)
            
        else:
            # Párrafo normal
            para = doc.add_paragraph()
            para.style = doc.styles['Normal']
            
            # Procesar contenido con formato (negrita, cursiva)
            self._add_formatted_text(para, element, options)
    
    def _add_formatted_text(self, paragraph, element, options: ProfessionalDocumentOptions):
        """Agrega texto con formato (negrita, cursiva, etc.)."""
        
        for item in element.contents:
            if hasattr(item, 'name'):
                # Es un elemento HTML
                if item.name == 'strong' or item.name == 'b':
                    run = paragraph.add_run(item.get_text())
                    run.bold = True
                elif item.name == 'em' or item.name == 'i':
                    run = paragraph.add_run(item.get_text())
                    run.italic = True
                else:
                    run = paragraph.add_run(item.get_text())
            else:
                # Es texto plano
                paragraph.add_run(str(item))
    
    def _convert_to_pdf(self, word_path: str) -> str:
        """Convierte documento Word a PDF usando LibreOffice."""
        
        pdf_path = word_path.replace('.docx', '.pdf')
        
        try:
            # Intentar conversión con LibreOffice
            cmd = [
                'libreoffice',
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', os.path.dirname(word_path),
                word_path
            ]
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
            
            if result.returncode == 0 and os.path.exists(pdf_path):
                return pdf_path
            else:
                print(f"Error convertiendo a PDF: {result.stderr}")
                return None
                
        except (subprocess.TimeoutExpired, FileNotFoundError):
            print("LibreOffice no disponible para conversión a PDF")
            return None
    
    def cleanup(self):
        """Limpia archivos temporales."""
        try:
            import shutil
            shutil.rmtree(self.temp_dir)
        except Exception as e:
            print(f"Error limpiando archivos temporales: {e}")


def generate_professional_document(content: str, 
                                 book_data: Dict[str, Any],
                                 user_options: Dict[str, Any] = None) -> Dict[str, str]:
    """
    Función principal para generar documento profesional.
    
    Args:
        content: Contenido HTML del libro
        book_data: Información del libro (título, autor, etc.)
        user_options: Opciones de formateo del usuario
    
    Returns:
        Dict con paths a los documentos generados
    """
    
    # Crear opciones profesionales
    options = ProfessionalDocumentOptions()
    
    if user_options:
        # Aplicar configuraciones del usuario
        if 'font_family' in user_options:
            options.font_family = user_options['font_family']
        if 'font_size_body' in user_options:
            options.font_size_body = int(user_options['font_size_body'])
        if 'line_spacing' in user_options:
            options.line_spacing = float(user_options['line_spacing'])
        # ... más configuraciones según necesidad
    
    # Generar documento
    generator = ProfessionalDocumentGenerator()
    
    try:
        result = generator.generate_professional_document(content, book_data, options)
        return result
    except Exception as e:
        generator.cleanup()
        raise e


# Función de utilidad para instalación de dependencias
def check_dependencies():
    """Verifica e informa sobre dependencias necesarias."""
    
    missing_deps = []
    
    try:
        import docx
    except ImportError:
        missing_deps.append("python-docx")
    
    try:
        result = subprocess.run(['libreoffice', '--version'], 
                              capture_output=True, text=True, timeout=5)
        if result.returncode != 0:
            missing_deps.append("libreoffice")
    except (subprocess.TimeoutExpired, FileNotFoundError):
        missing_deps.append("libreoffice")
    
    if missing_deps:
        print("⚠️  Dependencias faltantes para generación profesional:")
        for dep in missing_deps:
            if dep == "python-docx":
                print(f"   - {dep}: pip install python-docx")
            elif dep == "libreoffice":
                print(f"   - {dep}: sudo apt-get install libreoffice (Linux) o descargar desde libreoffice.org")
        
        return False
    
    print("✅ Todas las dependencias están disponibles")
    return True


if __name__ == "__main__":
    # Test de dependencias
    if check_dependencies():
        print("🚀 Sistema listo para generación de documentos profesionales")
    else:
        print("❌ Instalar dependencias antes de continuar")